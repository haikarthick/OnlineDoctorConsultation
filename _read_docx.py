import docx
doc = docx.Document(r'c:\Users\HP\Downloads\issueList.docx')
print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] {p.text.strip()}')
print()
print('=== TABLES ===')
for ti, t in enumerate(doc.tables):
    print(f'--- Table {ti} ---')
    for ri, r in enumerate(t.rows):
        cells = [c.text.strip() for c in r.cells]
        print(f'  Row {ri}: ' + ' | '.join(cells))
