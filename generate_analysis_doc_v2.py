"""
VetCare Platform - Feature Analysis & Implementation Plan (V2)
Updated with Global Livestock Platform Research
Generates a professional Word document with embedded diagrams
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F2F7FB")
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    return table

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x37, 0x84, 0xC4)
    return heading

def add_body_text(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    return p

def save_fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf

# ── Diagram Generators ──────────────────────────────────────────────────

def create_architecture_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_facecolor('white')
    ax.text(7, 9.7, 'VetCare Platform - Post-Integration Architecture (Global)', 
            ha='center', va='center', fontsize=16, fontweight='bold', color='#1F4E79')
    frontend = FancyBboxPatch((0.3, 5.5), 13.4, 4, boxstyle="round,pad=0.1",
                               facecolor='#E8F4FD', edgecolor='#2E75B6', linewidth=2)
    ax.add_patch(frontend)
    ax.text(7, 9.2, 'FRONTEND (React + Vite + TypeScript)', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1F4E79')
    existing_modules = [
        ('Marketplace\n(Livestock+Auction\n+Equipment)', '#4CAF50'),
        ('Feed Inventory\n(+ SmartFeed AI\n+ Ration Mgr)', '#FF9800'),
        ('AI Copilot\n(+ Livestock\n+ Decision AI)', '#9C27B0'),
        ('Animals\n(+ IoT Sensors\n+ ID Cards)', '#2196F3'),
    ]
    for i, (name, color) in enumerate(existing_modules):
        x = 1.0 + i * 3.3
        box = FancyBboxPatch((x, 7.2), 2.8, 1.7, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='white', alpha=0.85, linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + 1.4, 8.05, name, ha='center', va='center', fontsize=8.5, 
                fontweight='bold', color='white')
    ax.text(7, 6.9, 'Enhanced Existing Modules', ha='center', va='center',
            fontsize=9, fontstyle='italic', color='#555')
    new_modules = [
        ('Milk Recording\n& Dairy Dashboard', '#E91E63'),
        ('Community Forum\n+ Live Auctions', '#00BCD4'),
        ('Market Intelligence\n& Price Analytics', '#795548'),
        ('Knowledge Hub\n& Weather/Climate', '#607D8B'),
    ]
    for i, (name, color) in enumerate(new_modules):
        x = 1.0 + i * 3.3
        box = FancyBboxPatch((x, 5.7), 2.8, 1.0, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='white', alpha=0.85, linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + 1.4, 6.2, name, ha='center', va='center', fontsize=8.5,
                fontweight='bold', color='white')
    ax.text(7, 5.5, 'New Modules (from Global Research)', ha='center', va='center',
            fontsize=9, fontstyle='italic', color='#555')
    ax.annotate('', xy=(7, 4.6), xytext=(7, 5.5),
                arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2.5))
    backend = FancyBboxPatch((0.3, 3.0), 13.4, 1.7, boxstyle="round,pad=0.1",
                              facecolor='#FFF3E0', edgecolor='#E65100', linewidth=2)
    ax.add_patch(backend)
    ax.text(7, 4.35, 'BACKEND (Express + TypeScript + Raw pg.Pool SQL)', ha='center', 
            va='center', fontsize=12, fontweight='bold', color='#E65100')
    backend_items = ['52+ Services', '24+ Controllers', 'API /api/v1', 'Joi Validation', 
                     'JWT Auth + RBAC']
    for i, item in enumerate(backend_items):
        x = 1.2 + i * 2.6
        box = FancyBboxPatch((x, 3.15), 2.2, 0.6, boxstyle="round,pad=0.05",
                              facecolor='#FF9800', edgecolor='white', alpha=0.8)
        ax.add_patch(box)
        ax.text(x + 1.1, 3.45, item, ha='center', va='center', fontsize=8, 
                fontweight='bold', color='white')
    ax.annotate('', xy=(5, 1.8), xytext=(5, 3.0),
                arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2.5))
    ax.annotate('', xy=(10, 1.8), xytext=(10, 3.0),
                arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2.5))
    db = FancyBboxPatch((0.3, 0.3), 6.5, 1.5, boxstyle="round,pad=0.1",
                         facecolor='#E8F5E9', edgecolor='#2E7D32', linewidth=2)
    ax.add_patch(db)
    ax.text(3.55, 1.5, 'PostgreSQL Database', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#2E7D32')
    ax.text(3.55, 1.0, '44 existing + 12 new tables + 3 extended', ha='center', 
            va='center', fontsize=9, color='#555')
    ax.text(3.55, 0.65, 'Schema-separated: vetcare_dev / vetcare_prod', ha='center',
            va='center', fontsize=8, fontstyle='italic', color='#777')
    ext = FancyBboxPatch((7.2, 0.3), 6.5, 1.5, boxstyle="round,pad=0.1",
                          facecolor='#F3E5F5', edgecolor='#7B1FA2', linewidth=2)
    ax.add_patch(ext)
    ax.text(10.45, 1.5, 'External Integrations', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#7B1FA2')
    ax.text(10.45, 1.0, 'Maps | SMS/OTP | Storage | Weather API', ha='center',
            va='center', fontsize=9, color='#555')
    ax.text(10.45, 0.65, 'AI/ML Engine | WebSocket | Email | RFID/EID', ha='center',
            va='center', fontsize=8, fontstyle='italic', color='#777')
    return save_fig_to_bytes(fig)


def create_global_platforms_map_diagram():
    """Create a world map showing researched platforms by country."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_facecolor('white')
    
    ax.text(7, 7.7, 'Global Livestock Platform Research - By Country', 
            ha='center', va='center', fontsize=15, fontweight='bold', color='#1F4E79')
    
    platforms = [
        {'country': 'INDIA', 'color': '#E91E63', 'x': 0.3, 'y': 5.8,
         'platforms': ['animall.in', 'pashushala.com'],
         'focus': 'Marketplace, AI, Feed'},
        {'country': 'USA', 'color': '#2196F3', 'x': 0.3, 'y': 4.1,
         'platforms': ['cattlerange.com', 'breedr.co (US)', 'farmbrite.com', 'herdwatch.com (US)'],
         'focus': 'Market Data, Supply Chain, Farm Mgmt'},
        {'country': 'AUSTRALIA', 'color': '#4CAF50', 'x': 7.2, 'y': 5.8,
         'platforms': ['auctionsplus.com.au', 'stocklive.com.au'],
         'focus': 'Live Online Auctions, Video Sales'},
        {'country': 'EUROPE', 'color': '#FF9800', 'x': 7.2, 'y': 4.1,
         'platforms': ['agriaffaires (FR/EU)', 'cowmanager.com (NL)', 'connecterra.ai (NL)', 'agrivi.com (HR)'],
         'focus': 'IoT Sensors, AI Analytics, Traceability'},
        {'country': 'UK / IRELAND', 'color': '#9C27B0', 'x': 0.3, 'y': 2.4,
         'platforms': ['breedr.co (UK)', 'herdwatch.com (IE)'],
         'focus': 'Cattle Tracking, Herd Mgmt'},
        {'country': 'TURKEY', 'color': '#00BCD4', 'x': 7.2, 'y': 2.4,
         'platforms': ['milkingcloud.com'],
         'focus': 'Dairy IoT, Cloud Mgmt'},
        {'country': 'GLOBAL', 'color': '#795548', 'x': 3.75, 'y': 0.5,
         'platforms': ['allflex.global (Merck)', 'dairyglobal.net'],
         'focus': 'RFID/EID Tags, Industry News'},
    ]
    
    for p in platforms:
        box = FancyBboxPatch((p['x'], p['y']), 6.6, 1.5, boxstyle="round,pad=0.1",
                              facecolor='white', edgecolor=p['color'], linewidth=2)
        ax.add_patch(box)
        # Country header
        hdr = FancyBboxPatch((p['x'] + 0.1, p['y'] + 1.1), 2.5, 0.3, boxstyle="round,pad=0.03",
                              facecolor=p['color'], edgecolor='white')
        ax.add_patch(hdr)
        ax.text(p['x'] + 1.35, p['y'] + 1.25, p['country'], ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='white')
        # Platforms list
        for i, plat in enumerate(p['platforms']):
            ax.text(p['x'] + 0.3, p['y'] + 0.85 - i * 0.2, f'• {plat}', va='center',
                    fontsize=7, color='#333')
        # Focus
        ax.text(p['x'] + 3.3, p['y'] + 0.6, f'Focus: {p["focus"]}', ha='left', va='center',
                fontsize=7, color='#555', fontstyle='italic')
    
    return save_fig_to_bytes(fig)


def create_marketplace_flow_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.text(7, 7.7, 'Enhanced Livestock Marketplace - Data Flow', 
            ha='center', va='center', fontsize=15, fontweight='bold', color='#1F4E79')
    steps_sell = [
        ('1. Select Animal Type\nCow/Buffalo/Goat/Sheep', '#E91E63'),
        ('2. Choose Breed\nSahiwal/HF/Gir/Murrah...', '#9C27B0'),
        ('3. Livestock Details\nMilk/Lactation/Pregnancy', '#673AB7'),
        ('4. Pricing\nFixed/Auction/Negotiable', '#3F51B5'),
        ('5. Photos\nSide + Udder (min 2)', '#2196F3'),
        ('6. Location\nGPS or Manual Pin', '#00BCD4'),
        ('7. Review & Publish', '#009688'),
    ]
    ax.text(3.5, 7.2, 'SELL FLOW (Multi-Step)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#E91E63')
    for i, (text, color) in enumerate(steps_sell):
        y = 6.5 - i * 0.85
        box = FancyBboxPatch((0.5, y - 0.25), 5.8, 0.55, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor='white', alpha=0.85)
        ax.add_patch(box)
        ax.text(3.4, y, text, ha='center', va='center', fontsize=8, 
                fontweight='bold', color='white')
        if i < len(steps_sell) - 1:
            ax.annotate('', xy=(3.4, y - 0.3), xytext=(3.4, y - 0.55),
                        arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))
    ax.text(10.5, 7.2, 'BUY FLOW (Discovery)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#2E7D32')
    card = FancyBboxPatch((7.5, 4.5), 6, 2.5, boxstyle="round,pad=0.1",
                           facecolor='#F5F5F5', edgecolor='#2E7D32', linewidth=2)
    ax.add_patch(card)
    ax.text(10.5, 6.7, 'LIVESTOCK LISTING CARD', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#2E7D32')
    card_fields = [
        'Breed: Sahiwal  |  ID: SAH-1613',
        'Lactation: 3rd  |  Milk: 12 L/day',
        'Pregnant: Yes (5 months)',
        'Age: 4y 2m  |  Weight: 420 kg',
        'Price: $1,200  |  [HOT DEAL]',
        'Location: Karnal, Haryana',
        'Vax: Up-to-date | EID: RFID Tagged',
    ]
    for i, field in enumerate(card_fields):
        ax.text(10.5, 6.3 - i * 0.25, field, ha='center', va='center',
                fontsize=7.5, color='#333', fontfamily='monospace')
    filters = FancyBboxPatch((7.5, 3.2), 6, 1.1, boxstyle="round,pad=0.05",
                              facecolor='#E3F2FD', edgecolor='#1565C0', linewidth=1.5)
    ax.add_patch(filters)
    ax.text(10.5, 3.95, 'Search & Filters', ha='center', va='center',
            fontsize=9, fontweight='bold', color='#1565C0')
    ax.text(10.5, 3.55, 'Species | Breed | Milk Range | Price | Pregnancy | Radius',
            ha='center', va='center', fontsize=7.5, color='#333')
    mapbox = FancyBboxPatch((7.5, 2.0), 6, 1.0, boxstyle="round,pad=0.05",
                             facecolor='#E8F5E9', edgecolor='#2E7D32', linewidth=1.5)
    ax.add_patch(mapbox)
    ax.text(10.5, 2.7, 'Map View + Live Auction Stream', ha='center', va='center',
            fontsize=9, fontweight='bold', color='#2E7D32')
    ax.text(10.5, 2.3, '10km / 25km / 50km / 100km radius + online auctions', ha='center',
            va='center', fontsize=8, color='#555')
    integ = FancyBboxPatch((0.5, 0.3), 13, 1.2, boxstyle="round,pad=0.1",
                            facecolor='#FFF8E1', edgecolor='#F57F17', linewidth=2)
    ax.add_patch(integ)
    ax.text(7, 1.2, 'Cross-Module Integration', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#F57F17')
    ax.text(7, 0.75, 'Animal Profile -> Auto-fill  |  Medical Records -> Vax Badge  |  '
            'Milk Records -> Verified Yield  |  Breeding -> Calving  |  '
            'Weather -> Climate Risk  |  Market Intel -> Price Benchmark',
            ha='center', va='center', fontsize=7.5, color='#555')
    return save_fig_to_bytes(fig)


def create_phase_timeline_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.text(7, 8.7, 'Implementation Roadmap - 5 Phases (Updated with Global Features)', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#1F4E79')
    phases = [
        {
            'title': 'PHASE 1: Core Revenue Drivers',
            'color': '#E91E63', 'bg': '#FCE4EC',
            'items': [
                '1. Livestock Marketplace Enhancement',
                '   (Rich cards + multi-step sell + live auction)',
                '2. Milk Recording & Dairy Dashboard',
                '   (Per-animal tracking + lactation analytics)',
                '3. Community Forum (Pashu Chat)',
                '   (Thread-based + expert badges)',
            ],
            'priority': 'HIGH',
        },
        {
            'title': 'PHASE 2: Ecosystem Growth',
            'color': '#FF9800', 'bg': '#FFF3E0',
            'items': [
                '4. Smart Feed AI + Ration Management',
                '5. Premium/Hot Deal Listings + Live Auction',
                '6. Medicine E-Commerce (Upchar)',
                '7. Market Price Intelligence Dashboard',
            ],
            'priority': 'HIGH-MED',
        },
        {
            'title': 'PHASE 3: Platform Expansion',
            'color': '#2196F3', 'bg': '#E3F2FD',
            'items': [
                '8. Equipment Marketplace (E-com + Rental)',
                '9. Dairy Products Marketplace + Subscriptions',
                '10. Partner & Entrepreneur Network',
                '11. Video Knowledge Hub + Pasture Mapping',
            ],
            'priority': 'MEDIUM',
        },
        {
            'title': 'PHASE 4: Intelligence & IoT',
            'color': '#9C27B0', 'bg': '#F3E5F5',
            'items': [
                '12. IoT Ear Sensor Monitoring (CowManager-style)',
                '13. Weather & Drought Integration (Climate Dashboard)',
                '14. Conception-to-Carcass Supply Chain Tracking',
                '15. RFID/EID Tag Integration + Auto-Drafting',
            ],
            'priority': 'MEDIUM',
        },
        {
            'title': 'PHASE 5: Polish & Scale',
            'color': '#4CAF50', 'bg': '#E8F5E9',
            'items': [
                '16. Live Activity Indicators + Viewer Count',
                '17. Phone OTP Authentication',
                '18. Animal ID Card System (QR + Print)',
                '19. Tipping / Platform Support + Livestock Finance',
            ],
            'priority': 'LOW',
        },
    ]
    for i, phase in enumerate(phases):
        y = 7.8 - i * 1.5
        box = FancyBboxPatch((0.3, y - 0.5), 13.4, 1.3, boxstyle="round,pad=0.1",
                              facecolor=phase['bg'], edgecolor=phase['color'], linewidth=2)
        ax.add_patch(box)
        title_box = FancyBboxPatch((0.5, y + 0.45), 4.5, 0.3, boxstyle="round,pad=0.05",
                                    facecolor=phase['color'], edgecolor='white')
        ax.add_patch(title_box)
        ax.text(2.75, y + 0.6, phase['title'], ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='white')
        badge = FancyBboxPatch((12.0, y + 0.45), 1.5, 0.3, boxstyle="round,pad=0.05",
                                facecolor=phase['color'], edgecolor='white', alpha=0.8)
        ax.add_patch(badge)
        ax.text(12.75, y + 0.6, phase['priority'], ha='center', va='center',
                fontsize=7.5, fontweight='bold', color='white')
        mid = len(phase['items']) // 2 + len(phase['items']) % 2
        for j, item in enumerate(phase['items'][:mid]):
            ax.text(1.0, y + 0.15 - j * 0.22, item, va='center', fontsize=7, color='#333')
        for j, item in enumerate(phase['items'][mid:]):
            ax.text(7.5, y + 0.15 - j * 0.22, item, va='center', fontsize=7, color='#333')
        if i < len(phases) - 1:
            ax.annotate('', xy=(7, y - 0.6), xytext=(7, y - 0.85),
                        arrowprops=dict(arrowstyle='->', color='#333', lw=2))
    return save_fig_to_bytes(fig)


def create_database_schema_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 11))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 11)
    ax.axis('off')
    ax.text(7, 10.7, 'Database Schema - New & Extended Tables (Updated)', ha='center', va='center',
            fontsize=15, fontweight='bold', color='#1F4E79')
    def draw_table(ax, x, y, title, fields, color, width=3.2, field_height=0.18):
        total_h = 0.3 + len(fields) * field_height + 0.05
        box = FancyBboxPatch((x, y - total_h), width, total_h, boxstyle="round,pad=0.05",
                              facecolor='#FAFAFA', edgecolor=color, linewidth=2)
        ax.add_patch(box)
        hdr = FancyBboxPatch((x, y - 0.3), width, 0.3, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor=color)
        ax.add_patch(hdr)
        ax.text(x + width/2, y - 0.15, title, ha='center', va='center',
                fontsize=7.5, fontweight='bold', color='white')
        for i, field in enumerate(fields):
            fy = y - 0.42 - i * field_height
            ax.text(x + 0.1, fy, field, va='center', fontsize=5.5, 
                    color='#333', fontfamily='monospace')
        return total_h
    ax.text(7, 10.3, 'NEW TABLES (12)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#2E7D32')
    draw_table(ax, 0.1, 9.9, 'milk_records', [
        'id SERIAL PK', 'animal_id FK', 'user_id FK',
        'record_date DATE', 'session ENUM', 'qty_liters DECIMAL',
        'fat_pct DECIMAL', 'snf_pct DECIMAL', 'temp DECIMAL'
    ], '#E91E63', 3.2)
    draw_table(ax, 3.5, 9.9, 'lactation_cycles', [
        'id SERIAL PK', 'animal_id FK', 'calving_date DATE',
        'lactation_number INT', 'peak_yield DECIMAL',
        'dry_off_date DATE', 'status ENUM'
    ], '#E91E63', 3.2)
    draw_table(ax, 6.9, 9.9, 'forum_threads', [
        'id SERIAL PK', 'category_id FK', 'user_id FK',
        'title VARCHAR', 'body TEXT', 'images JSONB',
        'is_pinned BOOL', 'reply_count INT',
        'upvote_count INT', 'created_at TS'
    ], '#00BCD4', 3.2)
    draw_table(ax, 10.5, 9.9, 'forum_replies', [
        'id SERIAL PK', 'thread_id FK', 'user_id FK',
        'parent_reply_id FK', 'body TEXT',
        'upvote_count INT', 'is_expert BOOL'
    ], '#00BCD4', 3.2)
    draw_table(ax, 0.1, 7.5, 'market_price_data', [
        'id SERIAL PK', 'species VARCHAR',
        'breed VARCHAR', 'region VARCHAR',
        'avg_price DECIMAL', 'min/max DECIMAL',
        'period_type ENUM', 'data_date DATE'
    ], '#795548', 3.2)
    draw_table(ax, 3.5, 7.5, 'live_auctions', [
        'id SERIAL PK', 'listing_id FK',
        'start_time TS', 'end_time TS',
        'stream_url VARCHAR', 'status ENUM',
        'current_bid DECIMAL', 'bid_count INT'
    ], '#FF9800', 3.2)
    draw_table(ax, 6.9, 7.5, 'auction_bids', [
        'id SERIAL PK', 'auction_id FK',
        'user_id FK', 'bid_amount DECIMAL',
        'bid_time TS', 'is_winning BOOL'
    ], '#FF9800', 3.2)
    draw_table(ax, 10.5, 7.5, 'partner_profiles', [
        'id SERIAL PK', 'user_id FK',
        'partner_type ENUM', 'business_name',
        'services JSONB', 'rating DECIMAL',
        'is_verified BOOL', 'commission_rate'
    ], '#795548', 3.2)
    draw_table(ax, 0.1, 5.5, 'knowledge_videos', [
        'id SERIAL PK', 'title VARCHAR',
        'video_url VARCHAR', 'category VARCHAR',
        'target_roles JSONB', 'is_published BOOL'
    ], '#607D8B', 3.2)
    draw_table(ax, 3.5, 5.5, 'weather_alerts', [
        'id SERIAL PK', 'region VARCHAR',
        'alert_type ENUM', 'severity ENUM',
        'description TEXT', 'valid_from/to TS',
        'data_source VARCHAR'
    ], '#FF5722', 3.2)
    draw_table(ax, 6.9, 5.5, 'iot_sensor_data', [
        'id SERIAL PK', 'animal_id FK',
        'sensor_type ENUM', 'reading JSONB',
        'alert_generated BOOL', 'timestamp TS'
    ], '#8BC34A', 3.2)
    draw_table(ax, 10.5, 5.5, 'supply_chain_events', [
        'id SERIAL PK', 'animal_id FK',
        'event_type ENUM', 'location VARCHAR',
        'handler_id FK', 'metadata JSONB',
        'timestamp TS'
    ], '#3F51B5', 3.2)
    ax.text(7, 3.5, 'EXTENDED EXISTING TABLES (3)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#E65100')
    draw_table(ax, 0.3, 3.1, 'marketplace_listings (ADD)', [
        '+ breed, species VARCHAR',
        '+ lactation_number INT',
        '+ daily_milk_yield DECIMAL',
        '+ pregnancy_status VARCHAR',
        '+ animal_weight_kg DECIMAL',
        '+ listing_tier ENUM',
        '+ is_hot_deal BOOLEAN',
        '+ linked_animal_id FK',
        '+ auction_type ENUM',
        '+ eid_tag_number VARCHAR',
    ], '#E65100', 4.2)
    draw_table(ax, 4.8, 3.1, 'animals (ADD)', [
        '+ current_lactation_num INT',
        '+ last_calving_date DATE',
        '+ daily_milk_yield DECIMAL',
        '+ lactation_status ENUM',
        '+ eid_tag_number VARCHAR',
        '+ pasture_group_id FK',
    ], '#E65100', 4.2)
    draw_table(ax, 9.3, 3.1, 'users (ADD)', [
        '+ phone_number VARCHAR',
        '+ phone_verified BOOLEAN',
        '+ partner_id FK',
        '+ livestock_finance_status',
    ], '#E65100', 4.2)
    return save_fig_to_bytes(fig)


def create_permission_matrix_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.text(6, 6.7, 'Permission Matrix - New Features by Role (Updated)', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#1F4E79')
    permissions = [
        'milk_recording', 'dairy_dashboard', 'community_forum', 'knowledge_hub',
        'partner_network', 'medicine_store', 'equipment_store', 'dairy_products',
        'boost_listing', 'market_intelligence', 'live_auctions', 'weather_climate',
        'iot_monitoring'
    ]
    roles = ['pet_owner', 'farmer', 'veterinarian', 'admin']
    matrix = [
        [False, True, True, True],   # milk_recording
        [False, True, False, True],  # dairy_dashboard
        [True, True, True, True],    # community_forum
        [True, True, True, True],    # knowledge_hub
        [False, True, False, True],  # partner_network
        [True, True, True, True],    # medicine_store
        [False, True, False, True],  # equipment_store
        [True, True, False, True],   # dairy_products
        [True, True, False, False],  # boost_listing
        [False, True, True, True],   # market_intelligence
        [True, True, True, True],    # live_auctions
        [False, True, True, True],   # weather_climate
        [False, True, True, True],   # iot_monitoring
    ]
    cell_w = 2.0
    cell_h = 0.37
    start_x = 3.5
    start_y = 6.0
    for j, role in enumerate(roles):
        x = start_x + j * cell_w
        box = FancyBboxPatch((x, start_y), cell_w - 0.05, cell_h, boxstyle="round,pad=0.02",
                              facecolor='#1F4E79', edgecolor='white')
        ax.add_patch(box)
        ax.text(x + cell_w/2, start_y + cell_h/2, role.replace('_', ' ').title(),
                ha='center', va='center', fontsize=8, fontweight='bold', color='white')
    for i, perm in enumerate(permissions):
        y = start_y - (i + 1) * cell_h
        hdr = FancyBboxPatch((0.3, y), 3.1, cell_h - 0.02, boxstyle="round,pad=0.02",
                              facecolor='#F5F5F5', edgecolor='#DDD')
        ax.add_patch(hdr)
        ax.text(1.85, y + cell_h/2, perm, ha='center', va='center',
                fontsize=7, fontweight='bold', color='#333')
        for j in range(len(roles)):
            x = start_x + j * cell_w
            has = matrix[i][j]
            color = '#4CAF50' if has else '#FFCDD2'
            symbol = '\u2713' if has else '\u2717'
            sym_color = 'white' if has else '#C62828'
            cell = FancyBboxPatch((x, y), cell_w - 0.05, cell_h - 0.02,
                                   boxstyle="round,pad=0.02",
                                   facecolor=color, edgecolor='white', alpha=0.8)
            ax.add_patch(cell)
            ax.text(x + cell_w/2, y + cell_h/2, symbol, ha='center', va='center',
                    fontsize=10, fontweight='bold', color=sym_color)
    return save_fig_to_bytes(fig)


def create_integration_map_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.text(7, 9.7, 'Cross-Module Integration Map (Global Features)', ha='center', va='center',
            fontsize=15, fontweight='bold', color='#1F4E79')
    hub = plt.Circle((7, 5), 1.2, facecolor='#1F4E79', edgecolor='white', linewidth=3)
    ax.add_patch(hub)
    ax.text(7, 5.2, 'ENHANCED', ha='center', va='center', fontsize=10,
            fontweight='bold', color='white')
    ax.text(7, 4.8, 'MARKETPLACE', ha='center', va='center', fontsize=10,
            fontweight='bold', color='white')
    modules = [
        (7, 8.5, 'Milk Recording\n& Dairy', '#E91E63', 'Verified yield'),
        (11, 7.5, 'Animal\nProfiles', '#2196F3', 'Auto-fill listing'),
        (12.5, 5, 'Medical\nRecords', '#4CAF50', 'Vax badge'),
        (11, 2.5, 'Breeding\n& Genetics', '#9C27B0', 'Calving + lineage'),
        (7, 1.5, 'Financial\nAnalytics', '#FF9800', 'Sale revenue'),
        (3, 2.5, 'Feed\nInventory', '#795548', 'SmartFeed recs'),
        (1.5, 5, 'Geospatial\nAnalytics', '#00BCD4', 'Map + nearby'),
        (3, 7.5, 'AI Copilot\n+ Decision', '#607D8B', 'AI insights'),
        (5, 8.7, 'Weather\n& Climate', '#FF5722', 'Drought alerts'),
        (9, 8.7, 'Community\nForum', '#009688', 'Share + discuss'),
        (12.5, 3.7, 'Supply\nChain', '#3F51B5', 'Traceability QR'),
        (1.5, 3.7, 'Wallet +\nFinance', '#F44336', 'Boost + livestock loan'),
        (1.5, 6.3, 'IoT Sensors\n+ EID Tags', '#8BC34A', 'Auto-sync data'),
        (12.5, 6.3, 'Market Price\nIntelligence', '#FF7043', 'Price benchmarks'),
    ]
    for mx, my, name, color, desc in modules:
        circle = plt.Circle((mx, my), 0.7, facecolor=color, edgecolor='white',
                             linewidth=2, alpha=0.85)
        ax.add_patch(circle)
        ax.text(mx, my + 0.1, name, ha='center', va='center', fontsize=7,
                fontweight='bold', color='white')
        ax.plot([mx, 7], [my, 5], color=color, linewidth=1.5, alpha=0.4,
                linestyle='--')
        mid_x = (mx + 7) / 2
        mid_y = (my + 5) / 2
        ax.text(mid_x, mid_y, desc, ha='center', va='center', fontsize=5.5,
                color='#555', fontstyle='italic',
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                         edgecolor='#DDD', alpha=0.9))
    return save_fig_to_bytes(fig)


def create_innovation_comparison_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.text(7, 8.7, 'Innovation Differentiators - VetCare vs All Competitors', 
            ha='center', va='center', fontsize=14, fontweight='bold', color='#1F4E79')
    features = [
        'AI-Verified Milk Yield',
        'Auto-populated Listings from Profile',
        'Disease-Aware Marketplace',
        'Feed Impact Tracker',
        'Blockchain Dairy Traceability',
        'Video Call Before Purchase',
        'Genomic Lineage on Listing Cards',
        'IoT-Verified Weight + Ear Sensors',
        'Multi-language Forum (5 langs)',
        'Subscription Dairy Delivery',
        'AI Decision Support + Copilot',
        'Weather/Drought-Aware Alerts',
        'Livestock Finance Integration',
        'Conception-to-Carcass Tracking',
    ]
    bar_h = 0.35
    start_y = 8.0
    for i, feat in enumerate(features):
        y = start_y - i * 0.5
        vetcare = FancyBboxPatch((4.5, y), 3.5, bar_h, boxstyle="round,pad=0.02",
                                  facecolor='#2E7D32', edgecolor='white', alpha=0.85)
        ax.add_patch(vetcare)
        ax.text(6.25, y + bar_h/2, '\u2713 VetCare', ha='center', va='center',
                fontsize=7, fontweight='bold', color='white')
        ind = FancyBboxPatch((8.2, y), 1.5, bar_h, boxstyle="round,pad=0.02",
                              facecolor='#FFCDD2', edgecolor='#EF9A9A', alpha=0.7)
        ax.add_patch(ind)
        ax.text(8.95, y + bar_h/2, '\u2717 India', ha='center', va='center',
                fontsize=6.5, color='#C62828')
        aus = FancyBboxPatch((9.9, y), 1.5, bar_h, boxstyle="round,pad=0.02",
                              facecolor='#FFCDD2', edgecolor='#EF9A9A', alpha=0.7)
        ax.add_patch(aus)
        ax.text(10.65, y + bar_h/2, '\u2717 AU/US', ha='center', va='center',
                fontsize=6.5, color='#C62828')
        eu = FancyBboxPatch((11.6, y), 1.5, bar_h, boxstyle="round,pad=0.02",
                             facecolor='#FFCDD2', edgecolor='#EF9A9A', alpha=0.7)
        ax.add_patch(eu)
        ax.text(12.35, y + bar_h/2, '\u2717 EU', ha='center', va='center',
                fontsize=6.5, color='#C62828')
        ax.text(4.3, y + bar_h/2, feat, ha='right', va='center',
                fontsize=7.5, fontweight='bold', color='#333')
    ax.text(7, 0.4, 'All 14 innovations are EXCLUSIVE to VetCare - no single global competitor offers all of these together',
            ha='center', va='center', fontsize=9, fontstyle='italic', color='#1F4E79',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#E8F4FD', edgecolor='#2E75B6'))
    return save_fig_to_bytes(fig)


# ── Main Document Generator ─────────────────────────────────────────────

def generate_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('VetCare Platform')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = 'Calibri'
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Feature Analysis & Implementation Plan')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source.add_run('Global Competitive Research - 15+ Platforms Across 8 Countries')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'
    run.italic = True
    
    doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('March 2026  |  Version 2.0  |  Confidential')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run.font.name = 'Calibri'
    
    doc.add_page_break()
    
    # ── TABLE OF CONTENTS ──
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Research Sources Overview',
        '   2.1  India - animall.in & pashushala.com',
        '   2.2  USA - CattleRange, Breedr, Farmbrite, Herdwatch',
        '   2.3  Australia - AuctionsPlus & StockLive',
        '   2.4  Europe - Agriaffaires, CowManager, Connecterra, AGRIVI',
        '   2.5  Turkey - MilkingCloud',
        '   2.6  Global - Allflex/Merck, DairyGlobal',
        '3. Global Platform Feature Matrix',
        '4. Gap Analysis - VetCare vs Global Leaders',
        '5. Architecture Diagram (Post-Integration)',
        '6. New Feature Specifications (from Global Research)',
        '   6.1  Live Online Auction System',
        '   6.2  Market Price Intelligence Dashboard',
        '   6.3  IoT Ear Sensor Monitoring',
        '   6.4  Weather & Drought Integration',
        '   6.5  Conception-to-Carcass Supply Chain',
        '   6.6  Pasture Management & Farm Mapping',
        '   6.7  Livestock Finance Integration',
        '   6.8  RFID/EID Tag Integration',
        '   6.9  Feed Ration Management (Enhanced)',
        '   6.10 Farm Accounting & eCommerce',
        '7. Original 14 Features (from India Research)',
        '8. Database Schema Additions (Updated)',
        '9. Permission System Changes (Updated)',
        '10. Cross-Module Integration Map',
        '11. Innovation Differentiators (Global)',
        '12. Implementation Roadmap (5 Phases)',
        '13. Impact Summary',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '1. Executive Summary', level=1)
    
    add_body_text(doc,
        'This document presents a comprehensive competitive analysis of 15+ livestock and farm management '
        'platforms across 8 countries - India, USA, Australia, UK/Ireland, Netherlands, France/Europe, '
        'Turkey, and global providers. It identifies unique features from each platform and proposes their '
        'integration into VetCare with significant innovations that NO SINGLE competitor currently offers.')
    
    add_body_text(doc,
        'The VetCare platform already possesses a robust foundation with 50+ pages, 44 backend services, '
        '55+ routes, and 5 supported languages. The proposed enhancements will transform VetCare from a '
        'veterinary consultation platform into a comprehensive global livestock ecosystem platform - '
        'combining the best features from Indian marketplaces, Australian live auctions, European IoT '
        'monitoring, American farm management, and more.')
    
    doc.add_paragraph()
    add_body_text(doc, 'Key Metrics (Updated):', bold=True)
    
    add_styled_table(doc,
        ['Metric', 'Original (India Only)', 'Updated (Global)'],
        [
            ['Platforms researched', '2', '15+'],
            ['Countries covered', '1 (India)', '8+'],
            ['New features to implement', '14', '19'],
            ['New frontend pages', '8-10', '12-15'],
            ['Existing pages to enhance', '3-4', '6-8'],
            ['New backend services', '4', '8'],
            ['New database tables', '8', '12'],
            ['Extended database tables', '2', '3'],
            ['New permission keys', '9', '13'],
            ['Implementation phases', '4', '5'],
        ])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 2. RESEARCH SOURCES
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '2. Research Sources Overview', level=1)
    
    # Global map diagram
    map_img = create_global_platforms_map_diagram()
    doc.add_picture(map_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ── 2.1 INDIA ──
    add_heading_styled(doc, '2.1 India - animall.in & pashushala.com', level=2)
    
    add_heading_styled(doc, 'animall.in', level=3)
    add_body_text(doc,
        "India's largest livestock trading platform with 1 Crore+ (10 million+) farmers. "
        "Primarily mobile-first with Hindi-language focus.")
    features_animall = [
        'Livestock Buy/Sell Marketplace by location (cow, buffalo, heifer, bull)',
        'Multi-Step Sell Form with animal type, calving number, daily milk, price, photos (side + udder), GPS',
        'Pashu Chat - Community forum for farmers ("1 Crore+ farmers trusted")',
        'Video Tutorials ("How to buy from Animall", "How to sell")',
        'Prime Listings - Featured/boosted animal listings for faster sales',
        'Tipping System, Nearby Animals discovery, Live Buyer Count',
    ]
    for feat in features_animall:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'pashushala.com', level=3)
    add_body_text(doc,
        '"Most Trusted Livestock Marketplace" with detailed animal cards, AI integration, '
        'and multiple e-commerce verticals.')
    features_pashushala = [
        'Rich Livestock Cards - Breed-specific IDs (SAH-1613, HF-2721), price, lactation, milk yield, pregnancy, location',
        'PashuGuru.AI - AI-powered agricultural assistant',
        'Aahar (Feed), Upchar (Medicine), Upkaran (Equipment), Utpad (Dairy Products) - 4 e-commerce verticals',
        'SmartFeed - AI-based intelligent feed optimization',
        'Pashu Vet - "Anytime Vet Support" consultation service',
        'Partner Program - Rural entrepreneurship network',
        '"Hot Deal" + "Book Now" badges on premium listings',
    ]
    for feat in features_pashushala:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    # ── 2.2 USA ──
    add_heading_styled(doc, '2.2 USA - CattleRange, Breedr, Farmbrite, Herdwatch', level=2)
    
    add_heading_styled(doc, 'cattlerange.com (USA)', level=3)
    add_body_text(doc,
        'Comprehensive US cattle market intelligence platform with listings, market data, and analysis tools.')
    features_cattlerange = [
        'Cattle for Sale indexed by Breed, Cattle Class, and State',
        'Daily & Weekly Market Summaries with price indices (fed cattle, feeder cattle, cow-calf)',
        'Regional Cattle Auction Results by state with historical comparison',
        'Cattle Calculators - Bred Heifer & Cow Value Estimator based on market data',
        'Weekly Drought Monitor integration (map-based drought severity by county)',
        'USDA Monthly Reports - Cold Storage, Livestock Slaughter, Cattle on Feed data',
        'Production Sale Calendar - Upcoming auction events with calendar view',
        'Feeder & Stocker Cattle Marketings data with trend analysis',
        'Industry news categorization (MCOOL, Beef Imports, Processing, Wildfires)',
    ]
    for feat in features_cattlerange:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'breedr.co (UK/USA)', level=3)
    add_body_text(doc,
        '"Powering the Future of Beef" - The world\'s most ambitious cattle management app. '
        'Tracks animals from conception to carcass, enabling supply chain optimization.')
    features_breedr = [
        'Lifetime Cattle Tracking - Conception to carcass, every animal, every step',
        'Supply Chain Collaboration - Independent producers collaborate within a supply chain',
        'Premium Beef Programs - Exclusive calf buy-back, retained ownership, preferential pricing grids',
        'Exponential Performance Gains - 30% feed conversion efficiency, 16% prime beef increase, $8/cwt premium',
        'Operation Types - Seedstock, Cow-Calf, Calf Nursery, Feedyard, Packers & Retailers',
        'Smart Beef Supply Chain management with data-driven decision making',
        'Case Studies & Producer Community platform',
    ]
    for feat in features_breedr:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'farmbrite.com (USA)', level=3)
    add_body_text(doc,
        'All-in-one farm management software for multi-species and biodiverse farms. '
        'Trusted by 5,000+ farmers worldwide with 8,000+ app integrations.')
    features_farmbrite = [
        'Livestock Management - End-to-end animal records, wellness, yields, analytics, herd management',
        'Farm Accounting & Bookkeeping - Receipt capture, split transactions, invoices, bank import',
        'Crop Planning & Management - Season planning from seed to sale',
        'Farm Sales & eCommerce - Direct-to-consumer sales, POS, order management, fulfillment',
        'Climate & Weather - Historic weather trends, satellite imagery for soil/livestock wellness',
        'Farm Mapping - Customized farm map builder for compliance and planning',
        '100+ Pre-built Reports - Operations, financial, and compliance reporting',
        'Resource & Input Management - Equipment utilization tracking, maintenance schedules',
        'Team Task Management - Kanban boards, recurring tasks, automated reminders',
        '8,000+ Zapier Integrations + Custom API',
    ]
    for feat in features_farmbrite:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'herdwatch.com (Ireland/USA)', level=3)
    add_body_text(doc,
        'Livestock management software used by 22,000+ producers globally. '
        '3 solutions in one platform: Cattle, Sheep, Pasture.')
    features_herdwatch = [
        'Cattle Management - Calving records, treatments, weights, breeding events, sales/purchases',
        'Sheep Management - Lamb recording, flock health, breeding performance tracking',
        'Pasture Management - Ranch mapping, fertilizer tracking, grazing tasks',
        'Offline-first with sync - Works without internet, syncs when connected',
        'Bluetooth EID Reader Integration - Connect to electronic ID readers and scale heads',
        'Multi-device support (phone, tablet, laptop) with cloud sync',
        'Breed Association Integration - Pedigree registrations, market declarations',
        'Bulk data upload - Send records to support team for batch import',
        'Recently acquired VetDrive (veterinary practice management system)',
    ]
    for feat in features_herdwatch:
        add_bullet(doc, feat)
    
    doc.add_page_break()
    
    # ── 2.3 AUSTRALIA ──
    add_heading_styled(doc, '2.3 Australia - AuctionsPlus & StockLive', level=2)
    
    add_heading_styled(doc, 'auctionsplus.com.au (Australia)', level=3)
    add_body_text(doc,
        "Australia's largest online livestock auction platform. 497+ upcoming auctions, "
        "24,488 head offered weekly, covering cattle, sheep, goats, equine, and machinery.")
    features_auctionsplus = [
        'Live Online Auctions - Real-time bidding with Watch or Bid options per auction',
        'Commodities - Cattle, Sheep, Goats, Equine, Machinery all on one platform',
        'Feeder Cattle Listings - Dedicated section for feeder/stocker cattle',
        'Trade Now - Instant purchase lots (no auction needed)',
        'Livestock Finance (StockCo) - Integrated financing for livestock purchases',
        'Trade & Finance Calculator - Calculate financing and trading costs',
        'Livestock Assessors - Find accredited assessors by region',
        'AuctionsPlus Young Cattle Indicator (AYCI) - Custom market index at 505c/kg',
        'Weekly Market Reports - Detailed reports with clearance rates, value-over-reserve',
        'Weather integration (temperature, wind, rain, cloud) on homepage',
        'Featured Lots with detailed breeding descriptions and pedigree data',
        'Forward+ Sheep - Forward contract marketplace for sheep',
        'Biosecurity Awareness resources and compliance',
        'Mobile App (iOS & Android) for bidding on-the-go',
    ]
    for feat in features_auctionsplus:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'stocklive.com.au (Australia)', level=3)
    add_body_text(doc,
        '"When you can\'t get to the auction, StockLive takes you there." '
        'Live-streamed livestock auctions across Australia (formerly Elite Livestock Auctions).')
    features_stocklive = [
        'Live Video Auction Streaming - Watch real-time auction with video feed',
        'Online Bidding - Bid remotely in real-time during live auctions',
        'Auction Calendar - Upcoming auctions by date, state (NSW, QLD, VIC, etc.)',
        'Commercial & Seedstock categories',
        'Per-page auction results with breed details, weights, and lot descriptions',
    ]
    for feat in features_stocklive:
        add_bullet(doc, feat)
    
    doc.add_page_break()
    
    # ── 2.4 EUROPE ──
    add_heading_styled(doc, '2.4 Europe - Agriaffaires, CowManager, Connecterra, AGRIVI', level=2)
    
    add_heading_styled(doc, 'agriaffaires.co.uk (France/Europe)', level=3)
    add_body_text(doc,
        'Major European farm equipment marketplace present in 20+ countries with 20+ languages. '
        '208,684 classified ads with 15,000 new ads daily.')
    features_agriaffaires = [
        'Livestock Equipment - Milking parlour, milk tanks, livestock trailers, head locks, water tanks, mixers',
        'Animals & Feed section - Live animals, straw & forage, feed, board',
        'Price Observatory - Average price trends tracking per equipment/animal category',
        'Equipment Transport - Request quotation based on distance for purchases',
        'Professional Directory - Searchable directory of agricultural businesses',
        'Property Sales - Farm land and property marketplace',
        'Geolocation Search - "All equipment around me" with radius filter',
        'Multi-country, multi-language support (20+ each)',
    ]
    for feat in features_agriaffaires:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'cowmanager.com (Netherlands)', level=3)
    add_body_text(doc,
        'IoT-based cow monitoring system using ear sensors for lifetime monitoring '
        'from calf to cow. Headquartered in Netherlands, offices in USA.')
    features_cowmanager = [
        'Ear Sensor Technology - Measures ear temperature + eating + rumination + activity patterns',
        'Health Monitoring - Alerts up to 3 days before clinical illness signs, mastitis detection',
        'Fertility/Heat Detection - Know exactly when cow is in heat, optimal insemination timing',
        'Transition Monitoring - Track dry period to calving transition health',
        'Nutrition Monitoring - Feed management optimization, heat stress prevention',
        'Youngstock Monitor - Calf health management for future milk production',
        '"Find My Cow" - Locator/Flash tools to find specific cows in herd within seconds',
        'MultiView - Give staff, vets, nutritionists access to selected cow data',
        '"Sort My Cow" - Automated cow separation with sorting gates for health/heat alerts',
        'Auto-Drafting - Automated separation of cows with alerts from herd',
    ]
    for feat in features_cowmanager:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'connecterra.ai (Netherlands)', level=3)
    add_body_text(doc,
        'AI-powered dairy analytics platform that removes data silos and provides '
        'AI-driven decision support for farms, advisors, and enterprises.')
    features_connecterra = [
        'Farm Data Analytics - Visualize, compare, analyze all farm data in one platform',
        'AI Copilot - Weekly operational summaries powered by AI, delivered to inbox',
        'Decision Support - Automatic impact analysis with AI-powered predictions',
        'Scenario Modeling - Model different decision scenarios with AI predictions',
        'Farm Timeline - Auto-track decisions and their outcomes over time',
        'ROI Calculator - Calculate impact of every decision on key metrics',
        'Multi-stakeholder - Solutions for farm leadership, dairy advisors, and enterprises',
        'Data Integration - Connects with all existing farm management systems',
    ]
    for feat in features_connecterra:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'agrivi.com (Croatia/Global)', level=3)
    add_body_text(doc,
        'Comprehensive digital agriculture solutions supporting the entire agri-food value chain. '
        'Partners include Nestle, Kaufland. UN and FAO affiliated.')
    features_agrivi = [
        'AGRIVI AI Engage - White-labeled WhatsApp/Viber AI agronomic advisor for farmer engagement',
        'Farm Management Software (FMS) - Data-driven tools for agronomic and economic decisions',
        'Food Traceability - Complete food production traceability from seed to shelves, QR code scanning',
        'IoT Integration (Connect) - Weather stations, soil sensors, machinery data plug & play',
        'ROI Calculator - Free farm digitalization assessment tool',
        'Multi-Industry - Farms, Enterprise Farms, Cooperatives, Food & Beverages, Banks, Retail, Public Sector',
        'Supply Chain Intelligence - Agriculture supply chain management for contracted farmers',
    ]
    for feat in features_agrivi:
        add_bullet(doc, feat)
    
    doc.add_page_break()
    
    # ── 2.5 TURKEY ──
    add_heading_styled(doc, '2.5 Turkey - MilkingCloud', level=2)
    
    add_heading_styled(doc, 'milkingcloud.com (Turkey)', level=3)
    add_body_text(doc,
        'Cloud-based dairy and beef farm management platform with integrated IoT devices. '
        'Available in English, Spanish, French, Turkish and more.')
    features_milkingcloud = [
        'Solutions - Dairy Farm, Beef Cattle/Feedlot, Cow-Calf Operations, Enterprise',
        'Devices - Mastitis Detection Device (MastiPro), Calving Sensor (PartuSense), '
        'Heat Detection (M2Moo ear tag), Wash Quality Monitor (WashLog)',
        'Modules - Dairy/Beef Cow Management, Group Management, Milk Yield Records, '
        'Weight Monitoring, Feed Management, Milk Ration Management, Beef Ration Management, '
        'Health Management, Heat/Breeding Operations, Pregnancy/Gestation Monitoring',
        'Freelance Veterinarian App - Multiple farms on single screen for traveling vets',
        'Enterprise Solutions - Multi-farm corporate management capabilities',
    ]
    for feat in features_milkingcloud:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    # ── 2.6 GLOBAL ──
    add_heading_styled(doc, '2.6 Global - Allflex/Merck, DairyGlobal', level=2)
    
    add_heading_styled(doc, 'allflex.global (Merck - Global)', level=3)
    add_body_text(doc,
        '65 years of livestock identification experience. Global leader in RFID tags, '
        'visual tags, tissue sampling, and SenseHub monitoring systems.')
    features_allflex = [
        'Visual ID Tags - Various colors, print options for cattle, sheep, goats, zoo, wildlife',
        'Electronic ID (RFID) - Species-specific RFID tags for traceability programs',
        'Tissue Sampling - Ear biopsies for health status, genomic insights, disease eradication',
        'SenseHub Monitoring - Cow-Calf, Dairy, Feedlot, and Poultry monitoring systems',
        'DataFlow II - Data management and flow for livestock intelligence',
        'Milking Automation - Automated milking system integration',
        'Tag Applicators + Readers - Hardware ecosystem for livestock identification',
        'Compliance - Supports national traceability and biosecurity regulatory programs',
    ]
    for feat in features_allflex:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'dairyglobal.net (Netherlands)', level=3)
    add_body_text(doc,
        "Gateway to the global dairy industry - comprehensive dairy news, expert opinions, "
        "webinars, digital magazines, and outbreak monitoring.")
    features_dairyglobal = [
        'Outbreak Updates - Real-time FMD (Foot-and-Mouth) and HPAI tracking globally',
        'Global Dairy Market Prices - Live price data across global markets',
        'Expert Opinions - Authored insights from industry leaders',
        'Webinars & Digital Magazine - Educational content for dairy professionals',
        'Focus Topics - Hoof health, Cows & Climate, Best practices',
    ]
    for feat in features_dairyglobal:
        add_bullet(doc, feat)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 3. GLOBAL PLATFORM FEATURE MATRIX
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '3. Global Platform Feature Matrix', level=1)
    
    add_body_text(doc,
        'Comparison of key feature categories across all researched platforms. '
        'This reveals which features are most common globally and where VetCare can differentiate.')
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Feature Category', 'India', 'USA', 'Australia', 'Europe', 'VetCare Status'],
        [
            ['Livestock Marketplace', 'animall, pashushala', 'Breedr', 'AuctionsPlus', 'Agriaffaires', 'Existing (enhance)'],
            ['Live Online Auctions', '-', '-', 'AuctionsPlus, StockLive', '-', 'NEW'],
            ['Herd/Cattle Management', '-', 'Herdwatch, Farmbrite', '-', 'CowManager', 'Existing (enhance)'],
            ['Milk Recording/Dairy', 'pashushala (basic)', 'Farmbrite', '-', 'MilkingCloud', 'NEW'],
            ['AI Advisory/Copilot', 'PashuGuru.AI', 'Connecterra Copilot', '-', 'AGRIVI AI', 'Existing (enhance)'],
            ['IoT Sensor Monitoring', '-', '-', '-', 'CowManager, MilkingCloud', 'Existing (enhance)'],
            ['Feed/Ration Management', 'SmartFeed', 'Farmbrite', '-', 'MilkingCloud', 'Existing (enhance)'],
            ['Market Price Intelligence', '-', 'CattleRange', 'AuctionsPlus AYCI', '-', 'NEW'],
            ['Weather/Drought Monitor', '-', 'CattleRange, Farmbrite', '-', 'AGRIVI Connect', 'NEW'],
            ['Supply Chain Traceability', '-', 'Breedr', '-', 'AGRIVI Food', 'Existing (enhance)'],
            ['RFID/EID Integration', '-', 'Herdwatch', '-', 'Allflex, CowManager', 'NEW'],
            ['Farm eCommerce/Sales', '-', 'Farmbrite', '-', '-', 'NEW'],
            ['Livestock Finance', '-', '-', 'AuctionsPlus StockCo', '-', 'NEW'],
            ['Pasture Management', '-', 'Herdwatch, Farmbrite', '-', '-', 'NEW'],
            ['Community Forum', 'Pashu Chat', '-', '-', '-', 'NEW'],
            ['Farm Mapping', '-', 'Farmbrite', '-', '-', 'Existing (enhance)'],
            ['Farm Accounting', '-', 'Farmbrite', '-', 'AGRIVI', 'NEW'],
            ['Vet Practice Management', '-', '-', '-', 'MilkingCloud', 'Existing'],
            ['Multi-language', 'Hindi only', 'English only', 'English only', '20+ (Agriaffaires)', 'Existing (5 langs)'],
        ],
        col_widths=[1.8, 1.3, 1.3, 1.3, 1.2])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 4. GAP ANALYSIS
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '4. Gap Analysis - VetCare vs Global Leaders', level=1)
    
    add_heading_styled(doc, '4.1 Already Covered (Enhancement Only)', level=2)
    
    add_styled_table(doc,
        ['Feature', 'VetCare Equivalent', 'Enhancement Needed'],
        [
            ['Livestock Marketplace', 'Marketplace (7 categories, auctions)', 'Add livestock-specific fields + live auction'],
            ['Feed Management', 'Feed Inventory (tracking)', 'Add ration management + cost optimization'],
            ['AI Copilot', 'AI Copilot (4 tools)', 'Add decision support + weekly summaries'],
            ['Vet Consultation', 'Vet Hospitals + Consultations + Video', 'Add emergency 24/7 + vet app'],
            ['Animal Records', 'Animals module', 'Add RFID/EID + pasture groups'],
            ['Multi-language', '5 languages', 'Already ahead of competitors'],
            ['Financial Tracking', 'Financial Analytics', 'Add farm accounting + auto milk revenue'],
            ['IoT Sensors', 'IoT Sensors module', 'Add ear sensor + calving sensor support'],
            ['Geospatial', 'Geospatial Analytics', 'Add farm mapping + pasture management'],
            ['Supply Chain', 'Supply Chain module', 'Add conception-to-carcass tracking'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '4.2 Genuinely New Features Required (Global)', level=2)
    
    add_styled_table(doc,
        ['#', 'Feature', 'Source Platform(s)', 'Priority'],
        [
            ['1', 'Live Online Auction System', 'AuctionsPlus, StockLive (AU)', 'HIGH'],
            ['2', 'Market Price Intelligence Dashboard', 'CattleRange (US), AuctionsPlus', 'HIGH'],
            ['3', 'IoT Ear Sensor Monitoring (Health/Heat/Nutrition)', 'CowManager (NL)', 'HIGH'],
            ['4', 'Weather & Drought Integration', 'CattleRange, Farmbrite (US), AGRIVI', 'MEDIUM'],
            ['5', 'Conception-to-Carcass Supply Chain', 'Breedr (UK/US)', 'MEDIUM'],
            ['6', 'Pasture Management & Farm Mapping', 'Herdwatch, Farmbrite (US)', 'MEDIUM'],
            ['7', 'Livestock Finance Integration', 'AuctionsPlus StockCo (AU)', 'MEDIUM'],
            ['8', 'RFID/EID Tag Integration', 'Allflex (Global), Herdwatch', 'MEDIUM'],
            ['9', 'Feed Ration Management (Dairy + Beef)', 'MilkingCloud (TR)', 'MEDIUM'],
            ['10', 'Farm eCommerce & Direct Sales', 'Farmbrite (US)', 'LOW'],
        ],
        col_widths=[0.3, 3.0, 2.5, 0.8])
    
    add_body_text(doc,
        'Note: These 10 NEW features are IN ADDITION to the original 14 features identified from '
        'Indian platform research (animall.in + pashushala.com). See Section 7 for original features.',
        italic=True)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 5. ARCHITECTURE DIAGRAM
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '5. Architecture Diagram (Post-Integration)', level=1)
    
    add_body_text(doc,
        'Updated architecture showing all proposed features from global research integrated. '
        'Enhanced existing modules in top row, new modules from global research in second row.')
    
    arch_img = create_architecture_diagram()
    doc.add_picture(arch_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 6. NEW FEATURE SPECIFICATIONS (FROM GLOBAL RESEARCH)
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '6. New Feature Specifications (from Global Research)', level=1)
    
    add_body_text(doc,
        'These 10 features are uniquely inspired by global platforms beyond India. Each specification '
        'details what the competitor does, how VetCare will implement it, and what innovations VetCare '
        'adds that the original platform lacks.')
    
    doc.add_paragraph()
    
    # ── 6.1 Live Online Auction ──
    add_heading_styled(doc, '6.1 Live Online Auction System', level=2)
    add_body_text(doc, 'Source: AuctionsPlus (Australia), StockLive (Australia)  |  New module', italic=True)
    
    add_body_text(doc, 'What Competitors Do:', bold=True)
    add_bullet(doc, 'AuctionsPlus: 497+ live auctions weekly, 24,488 head offered, real-time bidding with "Watch" or "Bid" options')
    add_bullet(doc, 'StockLive: Live video-streamed auctions, remote bidding during physical sales')
    
    doc.add_paragraph()
    add_body_text(doc, 'VetCare Implementation:', bold=True)
    
    add_styled_table(doc,
        ['Component', 'Description'],
        [
            ['Auction Scheduler', 'Sellers/admins schedule auctions with date, time, lot list, reserve prices'],
            ['Live Bid Engine', 'Real-time WebSocket bidding (leverages existing SocketProvider)'],
            ['Video Stream', 'Optional live video stream of animals during auction (WebRTC/embed)'],
            ['Watch Mode', 'Users can watch auction progress without bidding (builds engagement)'],
            ['Auction Types', 'Timed auction (online-only), Live+Online (simulcast), Buy-Now (fixed price)'],
            ['Post-Sale Settlement', 'Automated invoicing, payment via Wallet, delivery coordination'],
            ['Market Indicators', 'Custom VetCare Cattle Index (like AYCI) based on auction results'],
        ])
    
    doc.add_paragraph()
    add_body_text(doc, 'VetCare Innovation (Beyond Competitors):', bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet(doc, 'AI-Verified Animal Data on auction lots (medical records, milk yield, vaccination history auto-attached)')
    add_bullet(doc, 'Video Call with Seller before bidding (not just video stream of sale)')
    add_bullet(doc, 'Multi-language auction interface (5 languages) - AuctionsPlus is English only')
    
    doc.add_paragraph()
    
    # ── 6.2 Market Price Intelligence ──
    add_heading_styled(doc, '6.2 Market Price Intelligence Dashboard', level=2)
    add_body_text(doc, 'Source: CattleRange (USA), AuctionsPlus AYCI (Australia)  |  New module', italic=True)
    
    add_body_text(doc, 'What Competitors Do:', bold=True)
    add_bullet(doc, 'CattleRange: Daily/weekly market summaries, price indices, USDA monthly reports, regional auction results')
    add_bullet(doc, 'AuctionsPlus: Weekly cattle market reports, value-over-reserve tracking, Young Cattle Indicator (AYCI)')
    add_bullet(doc, 'Agriaffaires: Price Observatory tracking average prices by equipment/category')
    
    doc.add_paragraph()
    add_body_text(doc, 'VetCare Implementation:', bold=True)
    market_items = [
        'Price Dashboard - Regional average prices by species, breed, age class, weight class',
        'Price History Charts - Daily/weekly/monthly trend lines, seasonal patterns, YoY comparison',
        'Market Reports - Auto-generated weekly summaries from VetCare platform transactions + external data',
        'Cattle Value Calculator - Input: breed, age, weight, lactation, location -> Estimated fair market value',
        'Price Alerts - Notify farmers when market price for their breed exceeds threshold',
        'Market Benchmarking - Compare your listing price vs regional average (helps sellers price competitively)',
    ]
    for item in market_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'VetCare Innovation:', bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet(doc, 'AI Price Prediction - ML model predicts price trends based on Season + Breed + Region + Supply data')
    add_bullet(doc, 'Embedded in Listing Flow - When seller creates listing, dashboard shows "Your price is X% above/below market"')
    
    doc.add_paragraph()
    
    # ── 6.3 IoT Ear Sensor Monitoring ──
    add_heading_styled(doc, '6.3 IoT Ear Sensor Monitoring', level=2)
    add_body_text(doc, 'Source: CowManager (Netherlands), MilkingCloud (Turkey), Allflex SenseHub  |  Enhancement to IoT Sensors', italic=True)
    
    add_body_text(doc, 'What Competitors Do:', bold=True)
    add_bullet(doc, 'CowManager: Ear sensor measures temperature + eating + rumination + activity => 3-day early illness alerts, heat detection, transition monitoring, youngstock monitoring, "Find My Cow", auto-drafting/sorting gates')
    add_bullet(doc, 'MilkingCloud: Mastitis detector (MastiPro), calving sensor (PartuSense), heat detection ear tag (M2Moo), wash quality monitor (WashLog)')
    add_bullet(doc, 'Allflex SenseHub: Cow-Calf, Dairy, Feedlot monitoring with RFID + health alerts')
    
    doc.add_paragraph()
    add_body_text(doc, 'VetCare Implementation:', bold=True)
    iot_items = [
        'Sensor Dashboard - Real-time display of eating, rumination, activity, and temperature per animal',
        'Health Alert Engine - 3-day early illness warning (rumination decline + temperature spike = flag)',
        'Heat Detection - Automated fertility alerts with optimal insemination window',
        'Calving Prediction - Sensor-based calving alert (PartuSense-style) with SMS/push notification',
        'Mastitis Early Warning - Abnormal milking patterns + temperature = mastitis risk score',
        '"Find My Cow" - GPS/locator within farm using sensor proximity data',
        'Transition Monitoring - Track dry period to calving health metrics',
    ]
    for item in iot_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'VetCare Innovation:', bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet(doc, 'IoT data feeds into Marketplace listings - Buyer sees "IoT-Verified: Health Score 95/100"')
    add_bullet(doc, 'Disease Prediction AI uses IoT data + regional outbreak data for predictive alerts')
    add_bullet(doc, 'No competitor combines IoT ear sensor data WITH marketplace AND vet consultation in one platform')
    
    doc.add_paragraph()
    
    # ── 6.4 Weather & Drought ──
    add_heading_styled(doc, '6.4 Weather & Drought Integration', level=2)
    add_body_text(doc, 'Source: CattleRange (drought monitor), Farmbrite (climate), AGRIVI (weather stations)  |  New feature', italic=True)
    
    weather_items = [
        'Farm Weather Dashboard - Current conditions, 7-day forecast, historical trends for farm location',
        'Drought Severity Map - Regional drought monitoring with impact on livestock feed availability',
        'Heat Stress Alerts - When temperature-humidity index exceeds threshold, alert about livestock risk',
        'Seasonal Feed Planning - Weather forecast integrated with Feed Inventory for proactive procurement',
        'Satellite Imagery - Pasture health assessment via NDVI (Normalized Difference Vegetation Index)',
        'Climate Risk Score - Rate farm locations for climate resilience (drought, flood, extreme heat risk)',
    ]
    for item in weather_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    # ── 6.5 Conception-to-Carcass ──
    add_heading_styled(doc, '6.5 Conception-to-Carcass Supply Chain', level=2)
    add_body_text(doc, 'Source: Breedr (UK/USA)  |  Enhancement to Supply Chain module', italic=True)
    
    supply_items = [
        'Lifetime Animal Timeline - Every event from conception to final processing in one view',
        'Supply Chain Roles - Seedstock > Cow-Calf > Nursery > Feedyard > Packer/Retailer chain',
        'Cross-Producer Collaboration - Independent producers share data within a supply chain',
        'Calf Buy-Back Programs - Seedstock sellers offer buy-back agreements to cow-calf operators',
        'Retained Ownership Models - Track animals through chain while maintaining breeder ownership',
        'Performance Feedback Loop - Carcass data flows back to breeder for genetic improvement',
        'Premium Pricing Grids - Data-driven preferential pricing for consistent quality producers',
    ]
    for item in supply_items:
        add_bullet(doc, item)
    
    add_body_text(doc, 'VetCare Innovation:', bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet(doc, 'Breedr tracks supply chain but has NO vet consultation, NO disease prediction, NO marketplace auction. VetCare combines all.')
    
    doc.add_paragraph()
    
    # ── 6.6 Pasture Management ──
    add_heading_styled(doc, '6.6 Pasture Management & Farm Mapping', level=2)
    add_body_text(doc, 'Source: Herdwatch (Pasture), Farmbrite (Farm Mapping)  |  Enhancement to Geospatial', italic=True)
    
    pasture_items = [
        'Interactive Farm Map Builder - Draw paddocks, buildings, water sources, lanes on satellite map',
        'Grazing Rotation Planner - Track which paddock each group is in, plan rotation schedule',
        'Fertilizer Application Tracking - Record what, when, how much applied per paddock',
        'Pasture Task Management - Create tasks for fencing, mowing, spraying with due dates',
        'Stocking Rate Calculator - Animals per hectare with carrying capacity recommendations',
    ]
    for item in pasture_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    # ── 6.7 Livestock Finance ──
    add_heading_styled(doc, '6.7 Livestock Finance Integration', level=2)
    add_body_text(doc, 'Source: AuctionsPlus StockCo (Australia)  |  New feature integrated with Marketplace', italic=True)
    
    finance_items = [
        'Livestock Purchase Financing - Buyer applies for financing during auction/purchase flow',
        'Trade & Finance Calculator - Estimate repayment schedule, interest, total cost before bidding',
        'Financing Status Tracking - Dashboard showing active loans, payments due, balance',
        'Partner Bank Integration - Connect with agricultural lenders/banks for loan processing',
        'Credit Scoring - Platform activity and transaction history build a livestock credit score',
    ]
    for item in finance_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    # ── 6.8 RFID/EID ──
    add_heading_styled(doc, '6.8 RFID/EID Tag Integration', level=2)
    add_body_text(doc, 'Source: Allflex (Global), Herdwatch (EID reader), CowManager  |  Enhancement to Animals module', italic=True)
    
    rfid_items = [
        'EID Tag Registration - Register RFID/EID tag numbers against animal profiles',
        'Bluetooth EID Reader Support - Scan tags via smartphone-connected reader',
        'Batch Scanning - Scan multiple animals for weighing, treatment recording, movement',
        'Traceability Compliance - Meet national livestock traceability regulatory requirements',
        'Automatic Animal Lookup - Scan tag -> instant animal profile with full history',
    ]
    for item in rfid_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    # ── 6.9 Feed Ration Management ──
    add_heading_styled(doc, '6.9 Feed Ration Management (Enhanced)', level=2)
    add_body_text(doc, 'Source: MilkingCloud (Turkey), Farmbrite (USA)  |  Enhancement to Feed Inventory', italic=True)
    
    ration_items = [
        'Dairy Ration Formulation - Create balanced rations for milk production optimization',
        'Beef Ration Management - Growth-optimized rations for feedlot cattle',
        'Nutritional Requirements Engine - Species + breed + weight + production stage = daily nutrient targets',
        'Ingredient Cost Optimizer - Least-cost ration that meets all nutritional requirements',
        'Ration Templates - Pre-built rations for common breed/stage combinations',
        'Feed Conversion Ratio Tracking - Measure efficiency (feed consumed vs output produced)',
    ]
    for item in ration_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    # ── 6.10 Farm eCommerce ──
    add_heading_styled(doc, '6.10 Farm eCommerce & Direct Sales', level=2)
    add_body_text(doc, 'Source: Farmbrite (USA)  |  Enhancement to Marketplace', italic=True)
    
    ecom_items = [
        'Farm Storefront - Each farmer gets a customizable online shop page',
        'Point-of-Sale (POS) - In-person sales recording at farm gate or market',
        'Order Management - Multi-channel order aggregation and fulfillment tracking',
        'Pick & Pack Lists - Automated picklists integrated with inventory',
        'Subscription Orders - Recurring delivery schedules for dairy products',
    ]
    for item in ecom_items:
        add_bullet(doc, item)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 7. ORIGINAL 14 FEATURES (from India Research)
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '7. Original 14 Features (from India Research)', level=1)
    
    add_body_text(doc,
        'These features were identified from the initial research of animall.in and pashushala.com. '
        'They remain part of the implementation plan and are detailed in the original analysis. '
        'Summary table below:')
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Priority'],
        [
            ['1', 'Livestock-Specific Listing Cards + Multi-Step Sell Flow', 'Both', 'HIGH'],
            ['2', 'Milk Recording & Dairy Dashboard', 'pashushala', 'HIGH'],
            ['3', 'Community Forum (Pashu Chat)', 'animall', 'HIGH'],
            ['4', 'Smart Feed AI Recommendations', 'pashushala', 'MEDIUM'],
            ['5', 'Veterinary Medicine E-Commerce (Upchar)', 'pashushala', 'MEDIUM'],
            ['6', 'Farm Equipment E-Commerce (Upkaran)', 'pashushala + agriaffaires', 'MEDIUM'],
            ['7', 'Dairy Products Marketplace (Utpad)', 'pashushala', 'MEDIUM'],
            ['8', 'Premium / Hot Deal Listings', 'Both', 'MEDIUM'],
            ['9', 'Partner & Entrepreneurship Network', 'pashushala', 'LOW'],
            ['10', 'Video Knowledge Hub (Tutorials)', 'animall', 'LOW'],
            ['11', 'Live Activity Indicators', 'animall', 'LOW'],
            ['12', 'Phone OTP Authentication', 'animall', 'LOW'],
            ['13', 'Animal ID Card System', 'pashushala', 'LOW'],
            ['14', 'Tipping / Platform Support', 'animall', 'LOW'],
        ],
        col_widths=[0.4, 3.5, 1.6, 0.8])
    
    add_body_text(doc,
        'Full specifications for these 14 features are in the original analysis document (V1). '
        'The combined total is 24 features across both phases of research.',
        italic=True)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 8. DATABASE SCHEMA
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '8. Database Schema Additions (Updated)', level=1)
    
    add_body_text(doc,
        'Updated schema showing 12 new tables and 3 extended tables to support both '
        'Indian and global feature sets.')
    
    db_img = create_database_schema_diagram()
    doc.add_picture(db_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '8.1 New Tables Summary (12)', level=2)
    
    add_styled_table(doc,
        ['Table', 'Purpose', 'Source Feature'],
        [
            ['milk_records', 'Per-animal, per-session milk recording', 'Milk Recording (India)'],
            ['lactation_cycles', 'Track lactation lifecycle per animal', 'Milk Recording (India)'],
            ['forum_threads', 'Community discussion threads', 'Community Forum (India)'],
            ['forum_replies', 'Thread replies with nesting', 'Community Forum (India)'],
            ['market_price_data', 'Regional market price intelligence', 'CattleRange (USA)'],
            ['live_auctions', 'Live auction events with streaming', 'AuctionsPlus (AU)'],
            ['auction_bids', 'Real-time bids for live auctions', 'AuctionsPlus (AU)'],
            ['partner_profiles', 'Partner registration and directory', 'pashushala (India)'],
            ['knowledge_videos', 'Video tutorials with categories', 'animall (India)'],
            ['weather_alerts', 'Regional weather and drought alerts', 'CattleRange/Farmbrite'],
            ['iot_sensor_data', 'IoT ear sensor readings and alerts', 'CowManager (NL)'],
            ['supply_chain_events', 'Conception-to-carcass event tracking', 'Breedr (UK/USA)'],
        ])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 9. PERMISSION SYSTEM
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '9. Permission System Changes (Updated)', level=1)
    
    add_body_text(doc,
        'Updated permission matrix with 13 new permission keys covering both Indian and global features. '
        'All 4 files must be updated in sync per VetCare architecture rules.')
    
    perm_img = create_permission_matrix_diagram()
    doc.add_picture(perm_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Permission Key', 'pet_owner', 'farmer', 'vet', 'admin', 'Source'],
        [
            ['milk_recording', '-', 'Y', 'Y', 'Y', 'India'],
            ['dairy_dashboard', '-', 'Y', '-', 'Y', 'India'],
            ['community_forum', 'Y', 'Y', 'Y', 'Y', 'India'],
            ['knowledge_hub', 'Y', 'Y', 'Y', 'Y', 'India'],
            ['partner_network', '-', 'Y', '-', 'Y', 'India'],
            ['medicine_store', 'Y', 'Y', 'Y', 'Y', 'India'],
            ['equipment_store', '-', 'Y', '-', 'Y', 'India'],
            ['dairy_products', 'Y', 'Y', '-', 'Y', 'India'],
            ['boost_listing', 'Y', 'Y', '-', '-', 'India'],
            ['market_intelligence', '-', 'Y', 'Y', 'Y', 'USA/AU'],
            ['live_auctions', 'Y', 'Y', 'Y', 'Y', 'Australia'],
            ['weather_climate', '-', 'Y', 'Y', 'Y', 'USA/EU'],
            ['iot_monitoring', '-', 'Y', 'Y', 'Y', 'Netherlands'],
        ],
        col_widths=[1.5, 0.8, 0.8, 0.6, 0.6, 1.0])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 10. INTEGRATION MAP
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '10. Cross-Module Integration Map', level=1)
    
    add_body_text(doc,
        'Updated integration map showing the enhanced Marketplace at center with connections to '
        '14 modules including new global features (Market Intelligence, Weather, IoT Sensors, Livestock Finance).')
    
    integ_img = create_integration_map_diagram()
    doc.add_picture(integ_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_body_text(doc, 'Key Integration Flows (Updated):', bold=True)
    flows = [
        'Animal Profile -> Auto-fill listing when marking animal "For Sale"',
        'Medical Records -> Vaccination badge on listing cards',
        'Milk Recording -> Verified daily yield on livestock listings (trust signal)',
        'Breeding Module -> Calving data + genetic lineage visible to buyers',
        'IoT Ear Sensors -> Real-time health score + "IoT-Verified" badge on listings',
        'Market Intelligence -> Price benchmark shown during listing creation + buyer comparison',
        'Weather/Climate -> Drought risk alerts on listings in affected regions',
        'Supply Chain -> Conception-to-carcass traceability QR on dairy products',
        'Live Auctions -> WebSocket real-time bidding integrated with marketplace',
        'Livestock Finance -> "Finance This Purchase" button on auction lots',
        'Community Forum -> Share listings for visibility + ask community about market prices',
        'AI Copilot -> Decision support with AI-powered weekly farm summaries',
        'Wallet -> Boost listing payments + auction escrow + finance repayments',
    ]
    for flow in flows:
        add_bullet(doc, flow)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 11. INNOVATION DIFFERENTIATORS
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '11. Innovation Differentiators (Global)', level=1)
    
    add_body_text(doc,
        'These 14 innovations make VetCare SIGNIFICANTLY SUPERIOR to ALL global competitors. '
        'No single platform worldwide combines all these capabilities.')
    
    innov_img = create_innovation_comparison_diagram()
    doc.add_picture(innov_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['#', 'Innovation', 'Why Unique'],
        [
            ['1', 'AI-Verified Milk Yield', 'ML validates claimed yield vs breed avg. No competitor verifies milk data.'],
            ['2', 'Auto-populated Listings', 'Seller selects animal -> listing pre-filled from profile + records. All others require manual entry.'],
            ['3', 'Disease-Aware Marketplace', 'Listings in outbreak zones get warning badges. Zero disease visibility on any competitor.'],
            ['4', 'Feed Impact Tracker', 'Change feed -> measure milk yield change over time. No feed-milk correlation anywhere.'],
            ['5', 'Blockchain Dairy Traceability', 'QR on dairy product -> full traceability. Neither India nor AU/US platforms offer this.'],
            ['6', 'Video Call Before Purchase', 'Buyer video-calls seller to see animal live. Not available on any platform.'],
            ['7', 'Genomic Lineage on Cards', 'Breeding quality score visible on listing. No genetic data shown on any marketplace.'],
            ['8', 'IoT-Verified Weight + Health', 'IoT ear sensor health score ON marketplace listing. CowManager monitors but has NO marketplace.'],
            ['9', 'Multi-language Forum', 'Auto-translate posts across 5 languages. Competitors are single-language only.'],
            ['10', 'Subscription Dairy', 'Full subscription management for daily delivery. No competitor has this.'],
            ['11', 'AI Decision Support + Copilot', 'Weekly AI summaries + scenario modeling (like Connecterra) PLUS vet consultation. Unique combo.'],
            ['12', 'Weather/Drought-Aware Alerts', 'CattleRange has drought map. VetCare integrates it WITH feed planning + listing alerts.'],
            ['13', 'Livestock Finance in Auction', 'AuctionsPlus has StockCo finance. VetCare adds it WITH AI credit scoring + vet records.'],
            ['14', 'Conception-to-Carcass + Vet', 'Breedr tracks supply chain. VetCare adds vet consultation + disease prediction at every stage.'],
        ],
        col_widths=[0.3, 2.0, 4.0])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 12. IMPLEMENTATION ROADMAP (5 Phases)
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '12. Implementation Roadmap (5 Phases)', level=1)
    
    phase_img = create_phase_timeline_diagram()
    doc.add_picture(phase_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Phase 1 - Core Revenue Drivers (HIGH)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Complexity'],
        [
            ['1', 'Livestock Marketplace Enhancement (rich cards + multi-step + live auction)', 'India + AU', 'HIGH'],
            ['2', 'Milk Recording & Dairy Dashboard', 'India + MilkingCloud', 'HIGH'],
            ['3', 'Community Forum (Pashu Chat)', 'India', 'HIGH'],
        ])
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Phase 2 - Ecosystem Growth (HIGH-MEDIUM)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Complexity'],
        [
            ['4', 'Smart Feed AI + Ration Management', 'India + MilkingCloud', 'MEDIUM'],
            ['5', 'Premium/Hot Deal + Live Auction System', 'India + AuctionsPlus', 'MEDIUM'],
            ['6', 'Medicine E-Commerce (Upchar)', 'India', 'MEDIUM'],
            ['7', 'Market Price Intelligence Dashboard', 'CattleRange + AuctionsPlus', 'MEDIUM'],
        ])
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Phase 3 - Platform Expansion (MEDIUM)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Complexity'],
        [
            ['8', 'Equipment Marketplace (E-commerce + Rental)', 'India + Agriaffaires', 'MEDIUM'],
            ['9', 'Dairy Products Marketplace + Subscriptions', 'India', 'MEDIUM'],
            ['10', 'Partner & Entrepreneur Network', 'India', 'MEDIUM'],
            ['11', 'Video Knowledge Hub + Pasture Management', 'India + Herdwatch/Farmbrite', 'LOW-MEDIUM'],
        ])
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Phase 4 - Intelligence & IoT (MEDIUM)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Complexity'],
        [
            ['12', 'IoT Ear Sensor Monitoring (CowManager-style)', 'CowManager (NL)', 'HIGH'],
            ['13', 'Weather & Drought Integration', 'CattleRange + Farmbrite', 'MEDIUM'],
            ['14', 'Conception-to-Carcass Supply Chain', 'Breedr (UK/USA)', 'HIGH'],
            ['15', 'RFID/EID Tag Integration + Auto-Drafting', 'Allflex + Herdwatch', 'MEDIUM'],
        ])
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Phase 5 - Polish & Scale (LOW)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Complexity'],
        [
            ['16', 'Live Activity Indicators + Viewer Count', 'animall', 'LOW'],
            ['17', 'Phone OTP Authentication', 'animall', 'MEDIUM'],
            ['18', 'Animal ID Card System (QR + Print)', 'pashushala', 'LOW'],
            ['19', 'Tipping / Platform Support + Livestock Finance', 'animall + AuctionsPlus', 'LOW-MEDIUM'],
        ])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════
    # 13. IMPACT SUMMARY
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, '13. Impact Summary', level=1)
    
    add_heading_styled(doc, 'Technical Impact', level=2)
    add_styled_table(doc,
        ['Metric', 'Before', 'After (All Phases)'],
        [
            ['Frontend Pages', '~50', '~65'],
            ['Backend Services', '44', '52+'],
            ['Database Tables', '~44', '~56'],
            ['API Routes', '55+', '80+'],
            ['Permission Keys', '~40', '~53'],
            ['Marketplace Categories', '7', '12+'],
            ['Nav Menu Items', '48', '60+'],
            ['IoT Device Types Supported', '3', '8+'],
            ['Languages', '5', '5 (ahead of all competitors)'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Business Impact', level=2)
    biz_impacts = [
        'Revenue Streams: Boost listings, marketplace commissions, auction fees, livestock finance commissions, partner commissions, subscription dairy, equipment rental commissions',
        'User Engagement: Community forum + milk recording create daily login habits; live auctions create event-based engagement',
        'Trust & Verification: AI-verified milk yield, IoT-verified health scores, vaccination badges, RFID traceability - unprecedented buyer confidence',
        'Market Intelligence: Price benchmarking helps sellers price competitively and buyers make informed decisions',
        'Ecosystem Lock-in: 14+ interconnected modules create deep data value impossible to replicate on any single competitor',
        'Geographic Expansion: 5-language support (EN, HI, TA, TE, KN) means immediate reach across India; global feature set attracts international users',
    ]
    for impact in biz_impacts:
        add_bullet(doc, impact)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Competitive Advantage Summary', level=2)
    add_body_text(doc,
        'VetCare will be the ONLY platform globally that combines: veterinary consultation, '
        'livestock marketplace with live auctions, dairy management, IoT sensor monitoring, '
        'AI-powered decision support, market price intelligence, weather/drought integration, '
        'conception-to-carcass supply chain, community forum, multi-vertical e-commerce '
        '(medicine + equipment + dairy products), livestock finance, and full traceability - '
        'all in a single integrated ecosystem with deep cross-module data flow in 5 languages.')
    
    doc.add_paragraph()
    add_body_text(doc,
        'No single competitor worldwide comes close to this level of integration:')
    comparison_items = [
        'animall.in/pashushala.com (India) - Marketplace only, no IoT, no auctions, no weather, single language',
        'AuctionsPlus (Australia) - Auctions only, no vet consultation, no IoT monitoring, no AI copilot',
        'CowManager (Netherlands) - IoT only, no marketplace, no community, no auctions',
        'Breedr (UK/USA) - Supply chain only, no vet consultation, no marketplace, no IoT dashboard',
        'Farmbrite (USA) - Farm management only, no auctions, no vet consultation, no IoT sensors',
        'Herdwatch (Ireland) - Herd management only, no marketplace, no AI copilot, no supply chain',
        'Connecterra (Netherlands) - AI analytics only, no marketplace, no auctions, no e-commerce',
        'MilkingCloud (Turkey) - Dairy IoT only, no marketplace, no community, no supply chain',
    ]
    for item in comparison_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run('--- End of Document ---')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'
    run.italic = True
    
    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p2.add_run('VetCare Platform  |  Global Feature Analysis & Implementation Plan  |  V2.0  |  March 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = 'Calibri'
    
    # ── Save ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                'VetCare_Feature_Analysis_Plan.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    path = generate_document()
    print(f"\nSuccess! Word document generated at:\n{path}")
    print(f"File size: {os.path.getsize(path) / 1024:.1f} KB")
