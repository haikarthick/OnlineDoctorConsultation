"""
VetCare Platform - Feature Analysis & Implementation Plan
Generates a professional Word document with embedded diagrams
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        set_cell_shading(cell, header_color)
    
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F2F7FB")
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    return table

def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x37, 0x84, 0xC4)
    return heading

def add_body_text(doc, text, bold=False, italic=False, color=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    return p

def save_fig_to_bytes(fig):
    """Convert matplotlib figure to bytes for embedding."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf

# ── Diagram Generators ──────────────────────────────────────────────────

def create_architecture_diagram():
    """Create high-level architecture diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_facecolor('white')
    
    # Title
    ax.text(7, 9.7, 'VetCare Platform - Post-Integration Architecture', 
            ha='center', va='center', fontsize=16, fontweight='bold', color='#1F4E79')
    
    # Frontend box
    frontend = FancyBboxPatch((0.3, 5.5), 13.4, 4, boxstyle="round,pad=0.1",
                               facecolor='#E8F4FD', edgecolor='#2E75B6', linewidth=2)
    ax.add_patch(frontend)
    ax.text(7, 9.2, 'FRONTEND (React + Vite + TypeScript)', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1F4E79')
    
    # Existing modules (enhanced)
    existing_modules = [
        ('Marketplace\n(Livestock Cards,\nAuctions, E-com)', '#4CAF50'),
        ('Feed Inventory\n(+ SmartFeed AI)', '#FF9800'),
        ('AI Copilot\n(+ Livestock\nPrompts)', '#9C27B0'),
        ('Animals\n(+ Lactation\nFields)', '#2196F3'),
    ]
    for i, (name, color) in enumerate(existing_modules):
        x = 1.0 + i * 3.3
        box = FancyBboxPatch((x, 7.2), 2.8, 1.7, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='white', alpha=0.85, linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + 1.4, 8.05, name, ha='center', va='center', fontsize=8.5, 
                fontweight='bold', color='white')
    
    ax.text(7, 6.9, 'Enhanced Existing Modules', ha='center', va='center',
            fontsize=9, fontstyle='italic', color='#555')
    
    # New modules
    new_modules = [
        ('Milk Recording\n& Dairy Dashboard', '#E91E63'),
        ('Community Forum\n(Pashu Chat)', '#00BCD4'),
        ('Partner\nNetwork', '#795548'),
        ('Knowledge\nHub', '#607D8B'),
    ]
    for i, (name, color) in enumerate(new_modules):
        x = 1.0 + i * 3.3
        box = FancyBboxPatch((x, 5.7), 2.8, 1.0, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='white', alpha=0.85, linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + 1.4, 6.2, name, ha='center', va='center', fontsize=8.5,
                fontweight='bold', color='white')
    ax.text(7, 5.5, 'New Modules', ha='center', va='center',
            fontsize=9, fontstyle='italic', color='#555')
    
    # Arrow Frontend -> Backend
    ax.annotate('', xy=(7, 4.6), xytext=(7, 5.5),
                arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2.5))
    
    # Backend box
    backend = FancyBboxPatch((0.3, 3.0), 13.4, 1.7, boxstyle="round,pad=0.1",
                              facecolor='#FFF3E0', edgecolor='#E65100', linewidth=2)
    ax.add_patch(backend)
    ax.text(7, 4.35, 'BACKEND (Express + TypeScript + Raw pg.Pool SQL)', ha='center', 
            va='center', fontsize=12, fontweight='bold', color='#E65100')
    
    backend_items = ['48+ Services', '22+ Controllers', 'API /api/v1', 'Joi Validation', 
                     'JWT Auth + RBAC']
    for i, item in enumerate(backend_items):
        x = 1.2 + i * 2.6
        box = FancyBboxPatch((x, 3.15), 2.2, 0.6, boxstyle="round,pad=0.05",
                              facecolor='#FF9800', edgecolor='white', alpha=0.8)
        ax.add_patch(box)
        ax.text(x + 1.1, 3.45, item, ha='center', va='center', fontsize=8, 
                fontweight='bold', color='white')
    
    # Arrow Backend -> Database
    ax.annotate('', xy=(5, 1.8), xytext=(5, 3.0),
                arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2.5))
    
    # Arrow Backend -> External
    ax.annotate('', xy=(10, 1.8), xytext=(10, 3.0),
                arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2.5))
    
    # Database box
    db = FancyBboxPatch((0.3, 0.3), 6.5, 1.5, boxstyle="round,pad=0.1",
                         facecolor='#E8F5E9', edgecolor='#2E7D32', linewidth=2)
    ax.add_patch(db)
    ax.text(3.55, 1.5, 'PostgreSQL Database', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#2E7D32')
    ax.text(3.55, 1.0, '44 existing + 8 new tables + 2 extended', ha='center', 
            va='center', fontsize=9, color='#555')
    ax.text(3.55, 0.65, 'Schema-separated: vetcare_dev / vetcare_prod', ha='center',
            va='center', fontsize=8, fontstyle='italic', color='#777')
    
    # External box
    ext = FancyBboxPatch((7.2, 0.3), 6.5, 1.5, boxstyle="round,pad=0.1",
                          facecolor='#F3E5F5', edgecolor='#7B1FA2', linewidth=2)
    ax.add_patch(ext)
    ax.text(10.45, 1.5, 'External Integrations', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#7B1FA2')
    ax.text(10.45, 1.0, 'Maps API | SMS/OTP | File Storage', ha='center',
            va='center', fontsize=9, color='#555')
    ax.text(10.45, 0.65, 'AI/ML Engine | WebSocket | Email', ha='center',
            va='center', fontsize=8, fontstyle='italic', color='#777')
    
    return save_fig_to_bytes(fig)


def create_marketplace_flow_diagram():
    """Create the enhanced marketplace data flow diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    
    ax.text(7, 7.7, 'Enhanced Livestock Marketplace - Data Flow', 
            ha='center', va='center', fontsize=15, fontweight='bold', color='#1F4E79')
    
    # Sell flow (left)
    steps_sell = [
        ('1. Select Animal Type\nCow/Buffalo/Goat/Sheep', '#E91E63'),
        ('2. Choose Breed\nSahiwal/HF/Gir/Murrah...', '#9C27B0'),
        ('3. Livestock Details\nMilk/Lactation/Pregnancy', '#673AB7'),
        ('4. Pricing\nFixed/Auction/Negotiable', '#3F51B5'),
        ('5. Photos\nSide + Udder (min 2)', '#2196F3'),
        ('6. Location\nGPS or Manual Pin', '#00BCD4'),
        ('7. Review & Publish', '#009688'),
    ]
    
    ax.text(3.5, 7.2, 'SELL FLOW (Multi-Step)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#E91E63')
    
    for i, (text, color) in enumerate(steps_sell):
        y = 6.5 - i * 0.85
        box = FancyBboxPatch((0.5, y - 0.25), 5.8, 0.55, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor='white', alpha=0.85)
        ax.add_patch(box)
        ax.text(3.4, y, text, ha='center', va='center', fontsize=8, 
                fontweight='bold', color='white')
        if i < len(steps_sell) - 1:
            ax.annotate('', xy=(3.4, y - 0.3), xytext=(3.4, y - 0.55),
                        arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))
    
    # Buy flow (right)
    ax.text(10.5, 7.2, 'BUY FLOW (Discovery)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#2E7D32')
    
    # Animal card
    card = FancyBboxPatch((7.5, 4.5), 6, 2.5, boxstyle="round,pad=0.1",
                           facecolor='#F5F5F5', edgecolor='#2E7D32', linewidth=2)
    ax.add_patch(card)
    ax.text(10.5, 6.7, 'LIVESTOCK LISTING CARD', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#2E7D32')
    
    card_fields = [
        'Breed: Sahiwal  |  ID: SAH-1613',
        'Lactation: 3rd  |  Milk: 12 L/day',
        'Pregnant: Yes (5 months)',
        'Age: 4y 2m  |  Weight: 420 kg',
        'Price: ₹85,000  |  [HOT DEAL]',
        'Location: Karnal, Haryana',
        '★ Vaccination: Up-to-date  ✓ Verified',
    ]
    for i, field in enumerate(card_fields):
        ax.text(10.5, 6.3 - i * 0.25, field, ha='center', va='center',
                fontsize=7.5, color='#333', fontfamily='monospace')
    
    # Filters
    filters = FancyBboxPatch((7.5, 3.2), 6, 1.1, boxstyle="round,pad=0.05",
                              facecolor='#E3F2FD', edgecolor='#1565C0', linewidth=1.5)
    ax.add_patch(filters)
    ax.text(10.5, 3.95, 'Search & Filters', ha='center', va='center',
            fontsize=9, fontweight='bold', color='#1565C0')
    ax.text(10.5, 3.55, 'Species | Breed | Milk Range | Price | Pregnancy | Radius',
            ha='center', va='center', fontsize=7.5, color='#333')
    
    # Map view
    mapbox = FancyBboxPatch((7.5, 2.0), 6, 1.0, boxstyle="round,pad=0.05",
                             facecolor='#E8F5E9', edgecolor='#2E7D32', linewidth=1.5)
    ax.add_patch(mapbox)
    ax.text(10.5, 2.7, 'Map View - Nearby Animals', ha='center', va='center',
            fontsize=9, fontweight='bold', color='#2E7D32')
    ax.text(10.5, 2.3, '10km / 25km / 50km / 100km radius', ha='center',
            va='center', fontsize=8, color='#555')
    
    # Bottom integration
    integ = FancyBboxPatch((0.5, 0.3), 13, 1.2, boxstyle="round,pad=0.1",
                            facecolor='#FFF8E1', edgecolor='#F57F17', linewidth=2)
    ax.add_patch(integ)
    ax.text(7, 1.2, 'Cross-Module Integration', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#F57F17')
    ax.text(7, 0.75, 'Animal Profile → Auto-fill Listing  |  Medical Records → Vaccination Badge  |  '
            'Milk Records → Verified Yield  |  Breeding → Calving Data  |  '
            'Geospatial → Map  |  Financial → Sale Revenue',
            ha='center', va='center', fontsize=7.5, color='#555')
    
    return save_fig_to_bytes(fig)


def create_phase_timeline_diagram():
    """Create implementation phases timeline diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 7)
    ax.axis('off')
    
    ax.text(7, 6.7, 'Implementation Roadmap - 4 Phases', ha='center', va='center',
            fontsize=15, fontweight='bold', color='#1F4E79')
    
    phases = [
        {
            'title': 'PHASE 1: Core Revenue Drivers',
            'color': '#E91E63',
            'bg': '#FCE4EC',
            'items': [
                '1. Livestock Marketplace Enhancement',
                '   (Rich animal cards + multi-step sell flow)',
                '2. Milk Recording & Dairy Dashboard',
                '   (Per-animal tracking + analytics)',
                '3. Community Forum (Pashu Chat)',
                '   (Thread-based + expert badges)',
            ],
            'priority': 'HIGH',
        },
        {
            'title': 'PHASE 2: Ecosystem Growth',
            'color': '#FF9800',
            'bg': '#FFF3E0',
            'items': [
                '4. Smart Feed AI Recommendations',
                '5. Premium/Hot Deal Listings',
                '6. Medicine E-Commerce (Upchar)',
            ],
            'priority': 'MEDIUM',
        },
        {
            'title': 'PHASE 3: Platform Expansion',
            'color': '#2196F3',
            'bg': '#E3F2FD',
            'items': [
                '7. Equipment E-Commerce (Upkaran)',
                '8. Dairy Products Marketplace',
                '9. Partner & Entrepreneur Network',
                '10. Video Knowledge Hub',
            ],
            'priority': 'MEDIUM',
        },
        {
            'title': 'PHASE 4: Polish & Scale',
            'color': '#4CAF50',
            'bg': '#E8F5E9',
            'items': [
                '11. Live Activity Indicators',
                '12. Phone OTP Authentication',
                '13. Animal ID Card System',
                '14. Tipping / Platform Support',
            ],
            'priority': 'LOW',
        },
    ]
    
    for i, phase in enumerate(phases):
        y = 5.8 - i * 1.5
        # Phase box
        box = FancyBboxPatch((0.3, y - 0.5), 13.4, 1.3, boxstyle="round,pad=0.1",
                              facecolor=phase['bg'], edgecolor=phase['color'], linewidth=2)
        ax.add_patch(box)
        
        # Phase title
        title_box = FancyBboxPatch((0.5, y + 0.45), 4, 0.3, boxstyle="round,pad=0.05",
                                    facecolor=phase['color'], edgecolor='white')
        ax.add_patch(title_box)
        ax.text(2.5, y + 0.6, phase['title'], ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')
        
        # Priority badge
        badge = FancyBboxPatch((12.0, y + 0.45), 1.5, 0.3, boxstyle="round,pad=0.05",
                                facecolor=phase['color'], edgecolor='white', alpha=0.8)
        ax.add_patch(badge)
        ax.text(12.75, y + 0.6, phase['priority'], ha='center', va='center',
                fontsize=8, fontweight='bold', color='white')
        
        # Items
        col1_items = phase['items'][:len(phase['items'])//2 + len(phase['items'])%2]
        col2_items = phase['items'][len(phase['items'])//2 + len(phase['items'])%2:]
        
        for j, item in enumerate(col1_items):
            ax.text(1.0, y + 0.15 - j * 0.22, item, va='center',
                    fontsize=7.5, color='#333')
        
        for j, item in enumerate(col2_items):
            ax.text(7.5, y + 0.15 - j * 0.22, item, va='center',
                    fontsize=7.5, color='#333')
        
        # Arrow between phases
        if i < len(phases) - 1:
            ax.annotate('', xy=(7, y - 0.6), xytext=(7, y - 0.85),
                        arrowprops=dict(arrowstyle='->', color='#333', lw=2))
    
    return save_fig_to_bytes(fig)


def create_database_schema_diagram():
    """Create database schema additions diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    
    ax.text(7, 8.7, 'Database Schema - New & Extended Tables', ha='center', va='center',
            fontsize=15, fontweight='bold', color='#1F4E79')
    
    def draw_table(ax, x, y, title, fields, color, width=3.8, field_height=0.2):
        total_h = 0.35 + len(fields) * field_height + 0.1
        # Table body
        box = FancyBboxPatch((x, y - total_h), width, total_h, boxstyle="round,pad=0.05",
                              facecolor='#FAFAFA', edgecolor=color, linewidth=2)
        ax.add_patch(box)
        # Header
        hdr = FancyBboxPatch((x, y - 0.35), width, 0.35, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor=color)
        ax.add_patch(hdr)
        ax.text(x + width/2, y - 0.175, title, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='white')
        # Fields
        for i, field in enumerate(fields):
            fy = y - 0.5 - i * field_height
            ax.text(x + 0.15, fy, field, va='center', fontsize=6.5, 
                    color='#333', fontfamily='monospace')
        return total_h
    
    # New tables
    ax.text(7, 8.3, 'NEW TABLES (8)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#2E7D32')
    
    draw_table(ax, 0.2, 7.9, 'milk_records', [
        'id SERIAL PK', 'animal_id FK → animals', 'user_id FK → users',
        'record_date DATE', 'session (AM/PM/Night)', 'quantity_liters DECIMAL',
        'fat_pct DECIMAL', 'snf_pct DECIMAL', 'temperature DECIMAL',
        'notes TEXT', 'created_at TIMESTAMP'
    ], '#E91E63', 3.2)
    
    draw_table(ax, 3.6, 7.9, 'lactation_cycles', [
        'id SERIAL PK', 'animal_id FK → animals', 'calving_date DATE',
        'lactation_number INT', 'peak_yield DECIMAL',
        'dry_off_date DATE', 'status (active/dry/done)',
        'created_at TIMESTAMP'
    ], '#E91E63', 3.2)
    
    draw_table(ax, 7.0, 7.9, 'forum_categories', [
        'id SERIAL PK', 'name VARCHAR', 'description TEXT',
        'icon VARCHAR', 'sort_order INT', 'is_active BOOLEAN'
    ], '#00BCD4', 3.2)
    
    draw_table(ax, 10.5, 7.9, 'forum_threads', [
        'id SERIAL PK', 'category_id FK', 'user_id FK',
        'title VARCHAR', 'body TEXT', 'images JSONB',
        'is_pinned BOOLEAN', 'is_locked BOOLEAN',
        'view_count INT', 'reply_count INT',
        'upvote_count INT', 'created_at, updated_at'
    ], '#00BCD4', 3.2)
    
    draw_table(ax, 0.2, 4.5, 'forum_replies', [
        'id SERIAL PK', 'thread_id FK', 'user_id FK',
        'parent_reply_id FK (nullable)', 'body TEXT',
        'upvote_count INT', 'is_expert_answer BOOLEAN',
        'created_at TIMESTAMP'
    ], '#00BCD4', 3.2)
    
    draw_table(ax, 3.6, 4.5, 'forum_votes', [
        'id SERIAL PK', 'user_id FK', 'thread_id FK (nullable)',
        'reply_id FK (nullable)', 'vote_type (up/down)',
        'UNIQUE(user_id, thread_id)', 'UNIQUE(user_id, reply_id)'
    ], '#00BCD4', 3.2)
    
    draw_table(ax, 7.0, 4.5, 'partner_profiles', [
        'id SERIAL PK', 'user_id FK', 'partner_type ENUM',
        'business_name VARCHAR', 'description TEXT',
        'services JSONB', 'location VARCHAR',
        'rating DECIMAL', 'is_verified BOOLEAN',
        'commission_rate DECIMAL', 'created_at'
    ], '#795548', 3.2)
    
    draw_table(ax, 10.5, 4.5, 'knowledge_videos', [
        'id SERIAL PK', 'title VARCHAR', 'description TEXT',
        'video_url VARCHAR', 'category VARCHAR',
        'target_roles JSONB', 'thumbnail_url VARCHAR',
        'sort_order INT', 'view_count INT',
        'is_published BOOLEAN', 'created_at'
    ], '#607D8B', 3.2)
    
    # Extended tables
    ax.text(7, 1.7, 'EXTENDED EXISTING TABLES (2)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#E65100')
    
    draw_table(ax, 0.5, 1.3, 'marketplace_listings (ADD COLUMNS)', [
        '+ breed VARCHAR', '+ species VARCHAR', '+ lactation_number INT',
        '+ daily_milk_yield_liters DECIMAL', '+ pregnancy_status VARCHAR',
        '+ pregnancy_months INT', '+ animal_age_months INT',
        '+ animal_weight_kg DECIMAL', '+ location_district VARCHAR',
        '+ location_state VARCHAR', '+ listing_tier (basic/silver/gold)',
        '+ is_hot_deal BOOLEAN', '+ linked_animal_id FK',
    ], '#E65100', 6.2)
    
    draw_table(ax, 7.3, 1.3, 'animals (ADD COLUMNS)', [
        '+ current_lactation_number INT',
        '+ last_calving_date DATE',
        '+ daily_milk_yield DECIMAL',
        '+ lactation_status (active/dry/n_a)',
    ], '#E65100', 6.2)
    
    return save_fig_to_bytes(fig)


def create_permission_matrix_diagram():
    """Create role-permission matrix diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(12, 6))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis('off')
    
    ax.text(6, 5.7, 'Permission Matrix - New Features by Role', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#1F4E79')
    
    permissions = [
        'milk_recording', 'dairy_dashboard', 'community_forum', 'knowledge_hub',
        'partner_network', 'medicine_store', 'equipment_store', 'dairy_products',
        'boost_listing'
    ]
    roles = ['pet_owner', 'farmer', 'veterinarian', 'admin']
    
    matrix = [
        [False, True, True, True],   # milk_recording
        [False, True, False, True],  # dairy_dashboard
        [True, True, True, True],    # community_forum
        [True, True, True, True],    # knowledge_hub
        [False, True, False, True],  # partner_network
        [True, True, True, True],    # medicine_store
        [False, True, False, True],  # equipment_store
        [True, True, False, True],   # dairy_products
        [True, True, False, False],  # boost_listing
    ]
    
    cell_w = 2.0
    cell_h = 0.42
    start_x = 3.5
    start_y = 5.0
    
    # Column headers
    for j, role in enumerate(roles):
        x = start_x + j * cell_w
        box = FancyBboxPatch((x, start_y), cell_w - 0.05, cell_h, boxstyle="round,pad=0.02",
                              facecolor='#1F4E79', edgecolor='white')
        ax.add_patch(box)
        ax.text(x + cell_w/2, start_y + cell_h/2, role.replace('_', ' ').title(),
                ha='center', va='center', fontsize=8, fontweight='bold', color='white')
    
    # Row headers + cells
    for i, perm in enumerate(permissions):
        y = start_y - (i + 1) * cell_h
        # Row header
        hdr = FancyBboxPatch((0.3, y), 3.1, cell_h - 0.02, boxstyle="round,pad=0.02",
                              facecolor='#F5F5F5', edgecolor='#DDD')
        ax.add_patch(hdr)
        ax.text(1.85, y + cell_h/2, perm, ha='center', va='center',
                fontsize=7.5, fontweight='bold', color='#333')
        
        for j in range(len(roles)):
            x = start_x + j * cell_w
            has = matrix[i][j]
            color = '#4CAF50' if has else '#FFCDD2'
            symbol = '✓' if has else '✗'
            sym_color = 'white' if has else '#C62828'
            
            cell = FancyBboxPatch((x, y), cell_w - 0.05, cell_h - 0.02, 
                                   boxstyle="round,pad=0.02",
                                   facecolor=color, edgecolor='white', alpha=0.8)
            ax.add_patch(cell)
            ax.text(x + cell_w/2, y + cell_h/2, symbol, ha='center', va='center',
                    fontsize=10, fontweight='bold', color=sym_color)
    
    return save_fig_to_bytes(fig)


def create_integration_map_diagram():
    """Create cross-module integration map."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    
    ax.text(7, 9.7, 'Cross-Module Integration Map', ha='center', va='center',
            fontsize=15, fontweight='bold', color='#1F4E79')
    
    # Central hub
    hub = plt.Circle((7, 5), 1.2, facecolor='#1F4E79', edgecolor='white', linewidth=3)
    ax.add_patch(hub)
    ax.text(7, 5.2, 'ENHANCED', ha='center', va='center', fontsize=10, 
            fontweight='bold', color='white')
    ax.text(7, 4.8, 'MARKETPLACE', ha='center', va='center', fontsize=10,
            fontweight='bold', color='white')
    
    # Surrounding modules
    modules = [
        (7, 8.5, 'Milk Recording\n& Dairy', '#E91E63', 
         'Verified milk yield\non listings'),
        (11, 7.5, 'Animal\nProfiles', '#2196F3',
         'Auto-fill listing\nfrom profile'),
        (12.5, 5, 'Medical\nRecords', '#4CAF50',
         'Vaccination badge\non listings'),
        (11, 2.5, 'Breeding\n& Genetics', '#9C27B0',
         'Calving data +\ngenetic profile'),
        (7, 1.5, 'Financial\nAnalytics', '#FF9800',
         'Sale revenue\nauto-tracking'),
        (3, 2.5, 'Feed\nInventory', '#795548',
         'SmartFeed recs +\nfeed e-commerce'),
        (1.5, 5, 'Geospatial\nAnalytics', '#00BCD4',
         'Map view +\nnearby animals'),
        (3, 7.5, 'AI Copilot', '#607D8B',
         'Livestock AI +\nfeed optimizer'),
        (5, 8.7, 'Wellness\nPortal', '#FF5722',
         'Health alerts on\nlisting cards'),
        (9, 8.7, 'Community\nForum', '#009688',
         'Share listings +\nask community'),
        (12.5, 3.7, 'Supply\nChain', '#3F51B5',
         'QR traceability\non dairy products'),
        (1.5, 3.7, 'Wallet', '#F44336',
         'Boost payments +\ntransactions'),
        (1.5, 6.3, 'IoT\nSensors', '#8BC34A',
         'Auto-sync weight\n& milk meters'),
        (12.5, 6.3, 'Disease\nPrediction', '#FF7043',
         'Outbreak warnings\non listings'),
    ]
    
    for mx, my, name, color, desc in modules:
        # Draw module circle
        circle = plt.Circle((mx, my), 0.7, facecolor=color, edgecolor='white', 
                             linewidth=2, alpha=0.85)
        ax.add_patch(circle)
        ax.text(mx, my + 0.1, name, ha='center', va='center', fontsize=7,
                fontweight='bold', color='white')
        
        # Draw connection line to hub
        ax.plot([mx, 7], [my, 5], color=color, linewidth=1.5, alpha=0.4, 
                linestyle='--')
        
        # Description label
        mid_x = (mx + 7) / 2
        mid_y = (my + 5) / 2
        ax.text(mid_x, mid_y, desc, ha='center', va='center', fontsize=5.5,
                color='#555', fontstyle='italic',
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white', 
                         edgecolor='#DDD', alpha=0.9))
    
    return save_fig_to_bytes(fig)


def create_innovation_comparison_diagram():
    """Create innovation differentiators comparison chart."""
    fig, ax = plt.subplots(1, 1, figsize=(13, 7))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.axis('off')
    
    ax.text(6.5, 6.7, 'Innovation Differentiators - VetCare vs Competitors', 
            ha='center', va='center', fontsize=14, fontweight='bold', color='#1F4E79')
    
    features = [
        'AI-Verified Milk Yield',
        'Auto-populated Listings',
        'Disease-Aware Marketplace',
        'Feed Impact Tracker',
        'Blockchain Dairy Traceability',
        'Video Call Before Purchase',
        'Genomic Lineage on Cards',
        'IoT-Verified Weight',
        'Multi-language Forum (5 langs)',
        'Subscription Dairy Delivery',
    ]
    
    bar_h = 0.4
    start_y = 6.0
    
    for i, feat in enumerate(features):
        y = start_y - i * 0.55
        
        # VetCare bar (full)
        vetcare = FancyBboxPatch((4.5, y), 4, bar_h, boxstyle="round,pad=0.02",
                                  facecolor='#2E7D32', edgecolor='white', alpha=0.85)
        ax.add_patch(vetcare)
        ax.text(6.5, y + bar_h/2, '✓ VetCare', ha='center', va='center',
                fontsize=8, fontweight='bold', color='white')
        
        # Competitor bars (empty/partial)
        animall = FancyBboxPatch((8.7, y), 2, bar_h, boxstyle="round,pad=0.02",
                                  facecolor='#FFCDD2', edgecolor='#EF9A9A', alpha=0.7)
        ax.add_patch(animall)
        ax.text(9.7, y + bar_h/2, '✗ animall', ha='center', va='center',
                fontsize=7, color='#C62828')
        
        pashushala = FancyBboxPatch((10.9, y), 2, bar_h, boxstyle="round,pad=0.02",
                                    facecolor='#FFCDD2', edgecolor='#EF9A9A', alpha=0.7)
        ax.add_patch(pashushala)
        ax.text(11.9, y + bar_h/2, '✗ pashushala', ha='center', va='center',
                fontsize=7, color='#C62828')
        
        # Feature name
        ax.text(4.3, y + bar_h/2, feat, ha='right', va='center',
                fontsize=8.5, fontweight='bold', color='#333')
    
    # Legend
    ax.text(6.5, 0.4, 'All 10 innovations are EXCLUSIVE to VetCare - neither competitor offers any of these',
            ha='center', va='center', fontsize=9, fontstyle='italic', color='#1F4E79',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#E8F4FD', edgecolor='#2E75B6'))
    
    return save_fig_to_bytes(fig)


# ── Main Document Generator ─────────────────────────────────────────────

def generate_document():
    doc = Document()
    
    # ── Page Setup ──
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('VetCare Platform')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = 'Calibri'
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Feature Analysis & Implementation Plan')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source.add_run('Based on Research of animall.in & pashushala.com')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'
    run.italic = True
    
    doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('March 2026  |  Confidential')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run.font.name = 'Calibri'
    
    # Page break
    doc.add_page_break()
    
    # ── TABLE OF CONTENTS ──
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Research Sources Overview',
        '3. Gap Analysis - Existing vs New Features',
        '4. Architecture Diagram (Post-Integration)',
        '5. Detailed Feature Specifications',
        '   5.1  Livestock Marketplace Enhancement',
        '   5.2  Milk Recording & Dairy Dashboard',
        '   5.3  Community Forum (Pashu Chat)',
        '   5.4  Smart Feed AI Recommendations',
        '   5.5  Veterinary Medicine E-Commerce',
        '   5.6  Farm Equipment E-Commerce',
        '   5.7  Dairy Products Marketplace',
        '   5.8  Premium/Hot Deal Listings',
        '   5.9  Partner & Entrepreneurship Network',
        '   5.10 Video Knowledge Hub',
        '   5.11 Live Activity & Engagement',
        '   5.12 Phone OTP Authentication',
        '   5.13 Animal ID Card System',
        '   5.14 Tipping / Platform Support',
        '6. Database Schema Additions',
        '7. Permission System Changes',
        '8. Cross-Module Integration Map',
        '9. Innovation Differentiators',
        '10. Implementation Roadmap',
        '11. Impact Summary',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '1. Executive Summary', level=1)
    
    add_body_text(doc, 
        'This document presents a comprehensive analysis of features from two leading Indian livestock '
        'platforms - animall.in and pashushala.com - and proposes their integration into the VetCare platform '
        'with significant improvements and innovations that neither competitor currently offers.')
    
    add_body_text(doc,
        'The VetCare platform already possesses a robust foundation with 50+ pages, 44 backend services, '
        '55+ routes, and 5 supported languages. The proposed enhancements will transform VetCare from a '
        'veterinary consultation platform into a comprehensive livestock ecosystem covering marketplace, '
        'dairy management, community, e-commerce, and partner networks.')
    
    doc.add_paragraph()
    add_body_text(doc, 'Key Metrics:', bold=True)
    
    add_styled_table(doc, 
        ['Metric', 'Value'],
        [
            ['New features to implement', '14'],
            ['New frontend pages', '8-10'],
            ['Existing pages to enhance', '3-4'],
            ['New backend services', '4'],
            ['New database tables', '8'],
            ['Extended database tables', '2'],
            ['New permission keys', '9'],
            ['Implementation phases', '4'],
        ])
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 2. RESEARCH SOURCES
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '2. Research Sources Overview', level=1)
    
    add_heading_styled(doc, '2.1 animall.in', level=2)
    add_body_text(doc, 
        'India\'s largest livestock trading platform with 1 Crore+ (10 million+) farmers. '
        'Primarily mobile-first with Hindi-language focus.')
    
    add_body_text(doc, 'Key Features Discovered:', bold=True)
    features_animall = [
        'Livestock Buy/Sell Marketplace - Browse cattle (cow, buffalo, heifer, bull) by location',
        'Multi-Step Sell Form - Animal type → calving number → daily milk → price → photos (side + udder) → GPS location',
        'Pashu Chat - Community forum for farmers (requires phone login), "1 Crore+ farmers trusted"',
        'Video Tutorials - "How to buy from Animall", "How to sell", "How to sell in 1 day"',
        'Prime Listings - Featured/boosted animal listings for faster sales',
        'Tipping System - Users can tip the platform',
        'Nearby Animals - Location-based discovery with radius filters',
        'Live Buyer Count - Shows real-time number of interested buyers',
        'Categories: Cow, Buffalo, Prime, Heifer, Young Female Buffalo, Male Buffalo, Bull, Other',
    ]
    for feat in features_animall:
        add_bullet(doc, feat)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '2.2 pashushala.com', level=2)
    add_body_text(doc,
        '"Most Trusted Livestock Marketplace" - More comprehensive than animall.in with detailed animal '
        'cards, AI integration, and multiple e-commerce verticals.')
    
    add_body_text(doc, 'Key Features Discovered:', bold=True)
    features_pashushala = [
        'Rich Livestock Cards - Breed-specific IDs (SAH-1613, HF-2721), price, lactation/biyat number, daily milk (L/day), pregnancy status+months, district/state, age, weight',
        'PashuGuru.AI - AI-powered agricultural assistant for livestock management',
        'Aahar (Feed/Nutrition) - Feed products marketplace and recommendations',
        'Upchar (Medicine/Treatment) - Veterinary medicine e-commerce',
        'Upkaran (Tools/Equipment) - Farm equipment marketplace',
        'Utpad (Products) - Dairy products marketplace (milk, ghee, butter)',
        'SmartFeed - AI-based intelligent feed optimization',
        'Pashu Vet - "Anytime Vet Support" veterinary consultation service',
        'Rural Entrepreneurship Network - Community for rural entrepreneurs',
        'Partner Program - "Grow Your Business - Join As Partner"',
        '"Hot Deal" + "Book Now" badges on premium listings',
        'Species: Cows (Sahiwal, HF, Gir, Jersey, Hariyana, Desi), Buffaloes (Murrah, Banni, Jafarabadi), Goats (Sirohi), Heifers',
    ]
    for feat in features_pashushala:
        add_bullet(doc, feat)
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 3. GAP ANALYSIS
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '3. Gap Analysis - Existing vs New Features', level=1)
    
    add_heading_styled(doc, '3.1 Already Covered (Needs Enhancement Only)', level=2)
    
    add_styled_table(doc,
        ['Feature (Competitor)', 'VetCare Equivalent', 'Enhancement Needed'],
        [
            ['Livestock Marketplace (buy/sell)', 'Marketplace (7 categories, auctions, orders)', 'Add livestock-specific fields (lactation, milk yield, pregnancy, calving)'],
            ['Aahar (Feed products)', 'Feed Inventory (tracking + consumption)', 'Add feed e-commerce layer'],
            ['PashuGuru.AI', 'AI Copilot (4 tools: chat, drugs, symptoms, scan)', 'Add livestock-specific prompts (feed optimization, breeding advice)'],
            ['Pashu Vet (Anytime Support)', 'Vet Hospitals + Consultations + Video', 'Add emergency quick-connect and 24/7 badge'],
            ['Breed Database', 'Animals with Gir, Sahiwal, Murrah, etc.', 'Add breed-specific attribute templates'],
            ['Multi-language', '5 languages (EN, HI, TA, TE, KN)', 'Already ahead of competitors - no change needed'],
            ['Financial Tracking', 'Financial Analytics (income/expense)', 'Add auto milk-sales revenue tracking'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '3.2 Genuinely New Features Required', level=2)
    
    add_styled_table(doc,
        ['#', 'Feature', 'Source', 'Priority'],
        [
            ['1', 'Livestock-Specific Listing Cards + Multi-Step Sell Flow', 'Both', 'HIGH'],
            ['2', 'Milk Recording & Dairy Dashboard', 'pashushala', 'HIGH'],
            ['3', 'Community Forum (Pashu Chat)', 'animall', 'HIGH'],
            ['4', 'Smart Feed AI Recommendations', 'pashushala', 'MEDIUM'],
            ['5', 'Veterinary Medicine E-Commerce (Upchar)', 'pashushala', 'MEDIUM'],
            ['6', 'Farm Equipment E-Commerce (Upkaran)', 'pashushala', 'MEDIUM'],
            ['7', 'Dairy Products Marketplace (Utpad)', 'pashushala', 'MEDIUM'],
            ['8', 'Premium / Hot Deal Listings', 'Both', 'MEDIUM'],
            ['9', 'Partner & Entrepreneurship Network', 'pashushala', 'LOW'],
            ['10', 'Video Knowledge Hub (Tutorials)', 'animall', 'LOW'],
            ['11', 'Live Activity Indicators', 'animall', 'LOW'],
            ['12', 'Phone Number Auth (OTP)', 'animall', 'LOW'],
            ['13', 'Animal ID Card System', 'pashushala', 'LOW'],
            ['14', 'Tipping / Platform Support', 'animall', 'LOW'],
        ],
        col_widths=[0.4, 3.5, 1.2, 0.9])
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 4. ARCHITECTURE DIAGRAM
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '4. Architecture Diagram (Post-Integration)', level=1)
    
    add_body_text(doc, 
        'The diagram below shows the VetCare platform architecture after all proposed features are '
        'integrated. Enhanced existing modules are shown in the top row, new modules in the second row.')
    
    arch_img = create_architecture_diagram()
    doc.add_picture(arch_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 5. DETAILED FEATURE SPECIFICATIONS
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '5. Detailed Feature Specifications', level=1)
    
    # ── 5.1 Livestock Marketplace ──
    add_heading_styled(doc, '5.1 Livestock Marketplace Enhancement', level=2)
    add_body_text(doc, 'Source: animall.in (buy/sell cattle) + pashushala.com (detailed animal cards)', italic=True)
    add_body_text(doc, 'Existing module to enhance: Marketplace', italic=True)
    
    doc.add_paragraph()
    add_body_text(doc, 'Current State:', bold=True)
    add_body_text(doc, 
        'Generic listings with title, description, price, category, images, auction bidding. '
        'No livestock-specific fields like lactation, milk yield, or pregnancy status.')
    
    doc.add_paragraph()
    add_body_text(doc, 'A) Enhanced Animal Listing Card (Buy Side)', bold=True)
    
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['Breed', 'Select (from breed DB)', 'Auto-populated: Sahiwal, HF, Gir, Jersey, Murrah, etc.'],
            ['Lactation / Calving Number', 'Integer', '1st, 2nd, 3rd... (biyat number)'],
            ['Daily Milk Yield', 'Decimal (L/day)', 'Morning + evening split, total daily'],
            ['Pregnancy Status', 'Select', 'Not pregnant / Pregnant (+ months input)'],
            ['Age', 'Years + Months', 'Auto-calculated from DOB if animal is registered'],
            ['Weight', 'Decimal (kg)', 'Auto-synced from IoT scale if available'],
            ['Price', 'Decimal + Currency', 'With "Negotiable" toggle'],
            ['Location', 'District, State', 'GPS pin + map preview'],
            ['Photos', 'Image uploads', 'Minimum 2: side view + udder/front view (max 8)'],
            ['Auto-generated Listing ID', 'String', 'Breed-coded: SAH-1613, HF-2721, MUR-0098'],
            ['Vaccination Status', 'Auto-pulled', 'From medical records (linked animal)'],
            ['Seller Verification', 'Badge', 'Verified farmer/enterprise badge'],
            ['Hot Deal / Prime', 'Badge', 'System-auto or paid promotion'],
        ])
    
    doc.add_paragraph()
    add_body_text(doc, 'B) Multi-Step Sell Flow', bold=True)
    
    # Sell flow diagram
    sell_img = create_marketplace_flow_diagram()
    doc.add_picture(sell_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    add_body_text(doc, 'C) Nearby Animals Discovery', bold=True)
    nearby_items = [
        'Map view with animal listing pins (leverages existing Geospatial module)',
        'Radius-based search: 10km, 25km, 50km, 100km, or all',
        'Filter by species, breed, milk yield range, price range, pregnancy status',
    ]
    for item in nearby_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'Use Cases:', bold=True)
    use_cases_mp = [
        'Farmer wants to sell a cow → Multi-step form auto-populates known data from existing animal profile → publishes rich listing card',
        'Buyer browsing by breed → Filters: species, breed, min/max milk yield, price range, pregnancy status, location radius',
        'Location-based discovery → Map view shows nearby animals for sale with price pins',
        'Cross-sell from existing animals → Farmer marks registered animal "For Sale" → listing auto-populated from profile + medical records',
    ]
    for i, uc in enumerate(use_cases_mp, 1):
        add_bullet(doc, f'UC{i}: {uc}')
    
    doc.add_paragraph()
    add_body_text(doc, 'Integration Points:', bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_body_text(doc, 
        'Animal Module (profile data) → Medical Records (vaccination badge) → Breeding (calving data) → '
        'Geospatial (map view) → Financial Analytics (sale as income) → Milk Recording (verified yield) → '
        'IoT Sensors (auto weight)')
    
    doc.add_page_break()
    
    # ── 5.2 Milk Recording ──
    add_heading_styled(doc, '5.2 Milk Recording & Dairy Dashboard', level=2)
    add_body_text(doc, 'Source: Both platforms (daily milk yield is the most important livestock metric)', italic=True)
    add_body_text(doc, 'New module - no existing equivalent', italic=True)
    
    doc.add_paragraph()
    add_body_text(doc, 'A) Milk Recording', bold=True)
    
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['Animal', 'Select (FK)', 'From registered animals (cow/buffalo/goat)'],
            ['Date', 'Date', 'Recording date (defaults to today)'],
            ['Session', 'Select', 'AM (Morning) / PM (Evening) / Night'],
            ['Quantity', 'Decimal', 'Liters of milk collected'],
            ['Fat %', 'Decimal', 'Fat content percentage'],
            ['SNF %', 'Decimal', 'Solid-not-fat percentage'],
            ['Temperature', 'Decimal', 'Milk temperature at collection'],
            ['Notes', 'Text', 'Optional quality/observation notes'],
        ])
    
    doc.add_paragraph()
    add_body_text(doc, 'Bulk Recording Mode:', bold=True)
    add_body_text(doc,
        'Table view where farmer enters milk quantities for the entire herd in one screen. '
        'Rows = animals, Columns = AM/PM/Night quantity. Quick-save all at once.')
    
    doc.add_paragraph()
    add_body_text(doc, 'B) Dairy Dashboard', bold=True)
    dashboard_items = [
        'Total daily production (farm-wide aggregate)',
        'Top producers (ranked by daily yield)',
        'Production trends (line chart: daily/weekly/monthly)',
        'Lactation curve visualization per animal',
        'Average fat% and SNF% trends',
        'Feed-to-milk conversion ratio (links to Feed Inventory)',
        'Revenue estimation: milk qty × configurable price/liter → links to Financial Analytics',
    ]
    for item in dashboard_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'C) Lactation Management', bold=True)
    lactation_items = [
        'Track full lactation cycle: calving date → peak → decline → dry period → next calving',
        'Lactation number auto-incremented per calving event (from Breeding module)',
        'Expected dry-off date prediction based on breed averages',
        'Alerts: declining production, expected heat/breeding window, dry-off approaching',
    ]
    for item in lactation_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'Use Cases:', bold=True)
    uc_milk = [
        'Morning milk recording → Farmer opens app, selects herd, bulk-enters milk quantities for each animal',
        'Identify low performers → Dashboard highlights animals below average, flags potential health issues',
        'Lactation tracking → Farmer sees which animals are peak vs dry, plans breeding accordingly',
        'Financial projection → Dashboard: "Today\'s milk = 450L × ₹45/L = ₹20,250"',
        'Buyer verification → When listing for sale, milk recording history = verifiable proof of yield',
    ]
    for i, uc in enumerate(uc_milk, 1):
        add_bullet(doc, f'UC{i}: {uc}')
    
    doc.add_paragraph()
    add_body_text(doc, 'Integration Points:', bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_body_text(doc,
        'Animal Module, Breeding (calving → new lactation), Feed Inventory (feed-to-milk ratio), '
        'Financial Analytics (auto-log milk sales), Wellness (milk decline = health alert), '
        'IoT Sensors (automated milk meters), Marketplace (verified yield on listings)')
    
    doc.add_page_break()
    
    # ── 5.3 Community Forum ──
    add_heading_styled(doc, '5.3 Community Forum (Pashu Chat)', level=2)
    add_body_text(doc, 'Source: animall.in "Pashu Chat" (1 Crore+ farmers)', italic=True)
    add_body_text(doc, 'New module', italic=True)
    
    doc.add_paragraph()
    add_body_text(doc, 'A) Forum Structure', bold=True)
    
    add_styled_table(doc,
        ['Component', 'Description'],
        [
            ['Categories', 'General Discussion, Cattle Management, Health & Disease, Feed & Nutrition, Breeding, Marketplace Tips, Dairy Management, Goat Farming, Success Stories'],
            ['Threads', 'Title + body + image attachments, pinnable, lockable'],
            ['Replies', 'Nested 1 level deep, upvote/downvote system'],
            ['Search', 'Full-text search across all threads and replies'],
            ['Sorting', 'Latest, Most Popular, Most Replied, Unanswered'],
        ])
    
    doc.add_paragraph()
    add_body_text(doc, 'B) Expert Badges', bold=True)
    badges = [
        'Veterinarian Badge - automatically assigned for vet role users',
        'Top Contributor - based on engagement score (posts, helpful votes)',
        'Verified Farmer - linked to enterprise verification',
        'Admin/Moderator - for content moderation team',
    ]
    for badge in badges:
        add_bullet(doc, badge)
    
    doc.add_paragraph()
    add_body_text(doc, 'C) Moderation System', bold=True)
    mod_items = [
        'Admin moderation queue for reported content',
        'User report mechanism (spam, inappropriate, misinformation)',
        'Auto-flag system (profanity filter)',
        'Content guidelines modal on first post',
    ]
    for item in mod_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'Use Cases:', bold=True)
    uc_forum = [
        '"My cow\'s milk dropped 30% in a week, what should I check?" → Gets vet-verified answer',
        'Experienced farmer shares feeding schedule that increased milk yield → Community upvotes',
        'Farmer shares premium bull listing → Cross-promotion from marketplace',
        'Disease Prediction flags outbreak risk → Auto-creates discussion thread for affected region',
    ]
    for i, uc in enumerate(uc_forum, 1):
        add_bullet(doc, f'UC{i}: {uc}')
    
    doc.add_page_break()
    
    # ── 5.4 Smart Feed AI ──
    add_heading_styled(doc, '5.4 Smart Feed AI Recommendations', level=2)
    add_body_text(doc, 'Source: pashushala.com "SmartFeed"', italic=True)
    add_body_text(doc, 'Enhancement to: Feed Inventory + AI Copilot', italic=True)
    
    doc.add_paragraph()
    smart_feed_items = [
        'Feed Recommendation Engine - Input: animal profile (species, breed, weight, age, lactation stage, pregnancy, daily milk, health conditions) → Output: daily diet plan (feed type → quantity → frequency → cost)',
        'Feed Cost Optimizer - Input: available inventory + market prices → Output: least-cost balanced ration meeting nutritional requirements',
        'Seasonal Adjustments - Summer/winter/monsoon-specific recommended changes',
        'Goal-based Optimization - Maximize milk yield, weight gain, or reproductive health',
        'Substitution Suggestions - When a feed is low-stock, suggest alternatives',
        'Feed Impact Tracker - Change feed → measure milk yield impact over 7/14/30 days',
    ]
    for item in smart_feed_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    add_body_text(doc, 'Use Cases:', bold=True)
    uc_feed = [
        'New farmer registers 10 cows → SmartFeed generates per-animal daily diet plan',
        'Feed Inventory shows low hay stock → SmartFeed suggests alternative ration using available feeds',
        'Farmer\'s Sahiwal producing 8L/day → SmartFeed recommends protein supplement to push to 10L/day',
        'SmartFeed identifies ₹2,000/month savings by switching brand B concentrate to equivalent brand A',
    ]
    for i, uc in enumerate(uc_feed, 1):
        add_bullet(doc, f'UC{i}: {uc}')
    
    doc.add_page_break()
    
    # ── 5.5-5.8 E-Commerce & Premium ──
    add_heading_styled(doc, '5.5 Veterinary Medicine E-Commerce (Upchar)', level=2)
    add_body_text(doc, 'Source: pashushala.com  |  Enhancement to: Marketplace (new sub-category)', italic=True)
    
    med_items = [
        'Dedicated medicine catalog: dewormers, antibiotics, supplements, vaccines, topical, hormonal',
        'Prescription-linked ordering: vet writes prescription → farmer clicks "Order Medicine" → shops',
        'Dosage calculator: species + weight → recommended dose',
        'Medicine reminders linked to Wellness Portal',
        'Verified pharmacy seller program with quality assurance',
    ]
    for item in med_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.6 Farm Equipment E-Commerce (Upkaran)', level=2)
    add_body_text(doc, 'Source: pashushala.com  |  Enhancement to: Marketplace (new sub-category)', italic=True)
    
    equip_items = [
        'Equipment catalog: milking machines, feed choppers, weighing scales, fencing, water troughs, vet instruments, transport cages',
        'New/Used toggle with condition grading',
        'Rental option: equipment sharing between farms (hourly/daily/weekly)',
        'IoT-integrated equipment links to IoT Sensors module for smart equipment',
    ]
    for item in equip_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.7 Dairy Products Marketplace (Utpad)', level=2)
    add_body_text(doc, 'Source: pashushala.com  |  Enhancement to: Marketplace (new sub-category)', italic=True)
    
    dairy_items = [
        'Product types: fresh milk, ghee, butter, paneer, curd, cheese, whey',
        'Quality grading: fat%, source breed, organic/non-organic certification',
        'Subscription model: daily/weekly milk delivery with auto-billing',
        'Bulk orders for B2B buyers (dairy cooperatives, restaurants)',
        'Full traceability: linked to Supply Chain (batch tracking farm → customer)',
    ]
    for item in dairy_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.8 Premium / Hot Deal Listing System', level=2)
    add_body_text(doc, 'Source: Both platforms  |  Enhancement to: Marketplace + Wallet', italic=True)
    
    add_styled_table(doc,
        ['Tier', 'Price', 'Benefits'],
        [
            ['Basic', 'Free', 'Standard listing, normal search position'],
            ['Silver', '₹99 / 7 days', '2x visibility, Silver badge, priority in search'],
            ['Gold', '₹299 / 15 days', '5x visibility, Gold badge, featured carousel, analytics dashboard'],
        ])
    
    premium_items = [
        '"Hot Deal" badge: auto-tagged by system when price is below market average for breed+lactation',
        'Seller analytics: views, inquiries, conversion rate for boosted vs normal listings',
        'Wallet integration: boost payments deducted from platform wallet balance',
        'Platform revenue model: boost purchases drive sustainable revenue',
    ]
    doc.add_paragraph()
    for item in premium_items:
        add_bullet(doc, item)
    
    doc.add_page_break()
    
    # ── 5.9-5.14 Remaining features ──
    add_heading_styled(doc, '5.9 Partner & Entrepreneurship Network', level=2)
    add_body_text(doc, 'Source: pashushala.com  |  New module', italic=True)
    
    partner_items = [
        'Partner types: feed suppliers, medicine shops, equipment dealers, transporters, AI technicians, dairy processors',
        'Partner registration + verification workflow',
        'Partner directory with ratings and reviews',
        'Partner dashboard: leads, orders, earnings, commission tracking',
        'Training/certification program for partners',
        'Auto-recommendation: SmartFeed recommends partner products → partner gets lead',
    ]
    for item in partner_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.10 Video Knowledge Hub', level=2)
    add_body_text(doc, 'Source: animall.in  |  New lightweight module', italic=True)
    
    video_items = [
        'Categorized tutorials: Getting Started, Animal Care, Breeding Tips, Feed Management, Disease Prevention, Farm Business, Technology How-To',
        'Embed YouTube or hosted videos with thumbnail grid',
        'Role-based content: farmer tutorials, vet training, pet owner guides',
        'Bookmark/favorites system',
        'Contextual "Related Videos" sidebar on relevant module pages',
        'Admin content management panel',
    ]
    for item in video_items:
        add_bullet(doc, item)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.11 Live Activity Indicators', level=2)
    add_body_text(doc, 'Source: animall.in  |  WebSocket enhancement', italic=True)
    add_bullet(doc, 'Real-time viewer count on marketplace listings ("23 people viewing")')
    add_bullet(doc, 'Online sellers indicator')
    add_bullet(doc, 'Live buyer interest count ("47 buyers looking for Sahiwal cows")')
    add_bullet(doc, 'Leverages existing WebSocket/SocketProvider infrastructure')
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.12 Phone OTP Authentication', level=2)
    add_body_text(doc, 'Source: animall.in  |  Auth module enhancement', italic=True)
    add_bullet(doc, 'Phone number registration/login alongside existing email auth')
    add_bullet(doc, 'SMS OTP verification via provider (Twilio/MSG91/similar)')
    add_bullet(doc, 'Critical for rural farmers who may not have email setup')
    add_bullet(doc, 'Dual auth: users can link both phone and email to same account')
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.13 Animal ID Card System', level=2)
    add_body_text(doc, 'Source: pashushala.com  |  Animal module enhancement', italic=True)
    add_bullet(doc, 'Auto-generated breed-coded IDs: SAH-1613 (Sahiwal), HF-2721 (HF), GIR-0045, MUR-3201')
    add_bullet(doc, 'Printable ID card with QR code linking to animal profile')
    add_bullet(doc, 'Useful for marketplace verification + traceability')
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.14 Tipping / Platform Support', level=2)
    add_body_text(doc, 'Source: animall.in  |  Wallet enhancement', italic=True)
    add_bullet(doc, 'Users can "tip" or support the VetCare platform via Wallet')
    add_bullet(doc, 'Optional microtransaction (₹10, ₹50, ₹100, custom)')
    add_bullet(doc, 'Supporter badge on profile')
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 6. DATABASE SCHEMA
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '6. Database Schema Additions', level=1)
    
    add_body_text(doc,
        'The diagram below shows all new tables and extensions to existing tables required '
        'for the full integration.')
    
    db_img = create_database_schema_diagram()
    doc.add_picture(db_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '6.1 New Tables (8)', level=2)
    
    new_tables = [
        ['milk_records', 'Per-animal, per-session milk recording with fat%, SNF%, temperature', 'animal_id → animals, user_id → users'],
        ['lactation_cycles', 'Track lactation lifecycle: calving → peak → dry → next calving', 'animal_id → animals'],
        ['forum_categories', 'Discussion categories (General, Cattle, Health, Feed, Breeding, etc.)', 'None'],
        ['forum_threads', 'Discussion threads with title, body, images, pin/lock, vote counts', 'category_id → forum_categories, user_id → users'],
        ['forum_replies', 'Thread replies with nesting (1 level), expert answer marking', 'thread_id → forum_threads, user_id → users'],
        ['forum_votes', 'Upvote/downvote on threads and replies (unique per user)', 'user_id → users, thread_id/reply_id FKs'],
        ['partner_profiles', 'Partner registration with type, services, rating, verification', 'user_id → users'],
        ['knowledge_videos', 'Video tutorials with categories, role targeting, publish control', 'None'],
    ]
    
    add_styled_table(doc,
        ['Table', 'Purpose', 'Foreign Keys'],
        new_tables)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, '6.2 Extended Existing Tables (2)', level=2)
    
    add_body_text(doc, 'marketplace_listings - New Columns:', bold=True)
    mp_cols = [
        'breed VARCHAR - Animal breed from breed database',
        'species VARCHAR - Cow, Buffalo, Goat, Sheep, etc.',
        'lactation_number INT - Calving/biyat number',
        'daily_milk_yield_liters DECIMAL - Average daily production',
        'pregnancy_status VARCHAR - Not pregnant / Pregnant',
        'pregnancy_months INT - Months pregnant (if applicable)',
        'animal_age_months INT - Age in months',
        'animal_weight_kg DECIMAL - Current weight',
        'location_district, location_state VARCHAR - Seller location',
        'listing_tier ENUM (basic/silver/gold) - Boost level',
        'is_hot_deal BOOLEAN - System-flagged competitive price',
        'linked_animal_id FK → animals - Links to registered animal',
    ]
    for col in mp_cols:
        add_bullet(doc, col)
    
    doc.add_paragraph()
    add_body_text(doc, 'animals - New Columns:', bold=True)
    animal_cols = [
        'current_lactation_number INT - Current lactation cycle number',
        'last_calving_date DATE - Most recent calving event',
        'daily_milk_yield DECIMAL - Latest recorded daily yield',
        'lactation_status ENUM (active/dry/not_applicable) - Current status',
    ]
    for col in animal_cols:
        add_bullet(doc, col)
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 7. PERMISSION SYSTEM
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '7. Permission System Changes (4-File Sync)', level=1)
    
    add_body_text(doc,
        'Per VetCare architecture rules, ALL 4 files must be updated in sync when adding new '
        'permissions. The table below shows the 9 new permission keys and their role assignments.')
    
    doc.add_paragraph()
    add_body_text(doc, 'Files to update:', bold=True)
    files_list = [
        'backend/src/services/PermissionService.ts - DEFAULT_ROLE_PERMISSIONS + PERMISSION_CATEGORIES',
        'frontend/src/context/PermissionContext.tsx - PERMISSION_ROUTE_MAP, ROUTE_PERMISSION_MAP, NAV_PERMISSION_MAP',
        'frontend/src/components/Navigation.tsx - menuItems array (roles + section)',
        'frontend/src/App.tsx - <RoleRoute path="..."> with matching path key',
    ]
    for f in files_list:
        add_bullet(doc, f)
    
    doc.add_paragraph()
    
    # Permission matrix diagram
    perm_img = create_permission_matrix_diagram()
    doc.add_picture(perm_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Permission Key', 'pet_owner', 'farmer', 'veterinarian', 'admin'],
        [
            ['milk_recording', '-', '✓', '✓', '✓'],
            ['dairy_dashboard', '-', '✓', '-', '✓'],
            ['community_forum', '✓', '✓', '✓', '✓'],
            ['knowledge_hub', '✓', '✓', '✓', '✓'],
            ['partner_network', '-', '✓', '-', '✓'],
            ['medicine_store', '✓', '✓', '✓', '✓'],
            ['equipment_store', '-', '✓', '-', '✓'],
            ['dairy_products', '✓', '✓', '-', '✓'],
            ['boost_listing', '✓', '✓', '-', '-'],
        ])
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 8. INTEGRATION MAP
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '8. Cross-Module Integration Map', level=1)
    
    add_body_text(doc,
        'The enhanced Marketplace sits at the center of the VetCare ecosystem, with deep '
        'bi-directional integrations to 14 existing and new modules. This creates a "flywheel effect" '
        'where each module adds value to the marketplace and vice versa.')
    
    integ_img = create_integration_map_diagram()
    doc.add_picture(integ_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_body_text(doc, 'Key Integration Flows:', bold=True)
    flows = [
        'Animal Profile → Auto-fill listing when seller marks animal "For Sale"',
        'Medical Records → Vaccination badge automatically shown on listing cards',
        'Milk Recording → Verified daily yield displayed on livestock listings (trust signal)',
        'Breeding Module → Calving data + genetic lineage visible to buyers',
        'Feed Inventory → SmartFeed AI recommendations + feed e-commerce',
        'Geospatial → Map-based nearby animal discovery',
        'Financial Analytics → Sale revenue auto-tracked + milk sales estimation',
        'Supply Chain → QR code traceability on dairy products (farm-to-customer)',
        'IoT Sensors → Auto-sync weight from smart scales + milk meter readings',
        'Disease Prediction → Outbreak warning badges on listings in affected zones',
        'AI Copilot → Livestock-specific prompts + feed optimization queries',
        'Community Forum → Share listings for community visibility + ask questions',
        'Wellness Portal → Health alerts visible on animal listing cards',
        'Wallet → Boost listing payments + marketplace transactions',
    ]
    for flow in flows:
        add_bullet(doc, flow)
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 9. INNOVATION DIFFERENTIATORS
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '9. Innovation Differentiators', level=1)
    
    add_body_text(doc,
        'These 10 innovations make VetCare SIGNIFICANTLY SUPERIOR to both animall.in and pashushala.com. '
        'Neither competitor currently offers any of these capabilities.')
    
    innov_img = create_innovation_comparison_diagram()
    doc.add_picture(innov_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['#', 'Innovation', 'Description', 'Competitor Status'],
        [
            ['1', 'AI-Verified Milk Yield', 'ML model validates claimed yield vs breed avg + lactation curve → trust score', 'Neither has verification'],
            ['2', 'Auto-populated Listings', 'Seller selects registered animal → listing pre-filled from profile + records', 'Both require manual entry'],
            ['3', 'Disease-Aware Marketplace', 'Listings in outbreak zones get warning badges; buyers alerted', 'Zero disease visibility'],
            ['4', 'Feed Impact Tracker', 'Change feed → measure milk yield impact over 7/14/30 days', 'No feed-milk correlation'],
            ['5', 'Blockchain Dairy Traceability', 'QR on milk/ghee packet → full traceability (farm, cow, feed, date)', 'Neither offers traceability'],
            ['6', 'Video Call Before Purchase', 'Buyer requests live video of animal before purchase decision', 'Neither offers video'],
            ['7', 'Genomic Lineage on Cards', 'Breeding pair quality score + genetic profile visible to buyers', 'No genetic data shown'],
            ['8', 'IoT-Verified Weight', 'Weight from IoT scale auto-synced (vs manual entry)', 'Neither integrates IoT'],
            ['9', 'Multi-language Forum', 'Auto-translate forum posts across 5 languages', 'animall: Hindi only'],
            ['10', 'Subscription Dairy', 'Full subscription management + delivery + auto-billing', 'Neither has subscription'],
        ])
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 10. IMPLEMENTATION ROADMAP
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '10. Implementation Roadmap', level=1)
    
    phase_img = create_phase_timeline_diagram()
    doc.add_picture(phase_img, width=Inches(6.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Phase 1 - Core Revenue Drivers (HIGH Priority)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'New Pages', 'Backend Changes', 'Complexity'],
        [
            ['1', 'Livestock Marketplace Enhancement', 'Enhance Marketplace.tsx + new MarketplaceSell.tsx multi-step form', 'Add 12+ livestock columns to marketplace_listings, enhance MarketplaceService', 'HIGH'],
            ['2', 'Milk Recording & Dairy Dashboard', 'New MilkRecording.tsx + DairyDashboard.tsx', 'New tables: milk_records, lactation_cycles. New MilkRecordingService.ts', 'HIGH'],
            ['3', 'Community Forum', 'New CommunityForum.tsx + ForumThread.tsx', 'New tables: forum_categories/threads/replies/votes. New ForumService.ts', 'HIGH'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Phase 2 - Ecosystem Growth (MEDIUM Priority)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Changes', 'Complexity'],
        [
            ['4', 'Smart Feed AI Recommendations', 'Enhance Feed Inventory page + AI Copilot feed prompts', 'MEDIUM'],
            ['5', 'Premium/Hot Deal Listings', 'Enhance Marketplace + Wallet boost payment integration', 'LOW-MEDIUM'],
            ['6', 'Medicine E-Commerce (Upchar)', 'New marketplace sub-section + prescription linking', 'MEDIUM'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Phase 3 - Platform Expansion (MEDIUM Priority)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Changes', 'Complexity'],
        [
            ['7', 'Equipment E-Commerce (Upkaran)', 'New marketplace sub-section + rental system', 'MEDIUM'],
            ['8', 'Dairy Products Marketplace', 'New marketplace sub-section + subscription model', 'MEDIUM'],
            ['9', 'Partner & Entrepreneur Network', 'New PartnerNetwork.tsx + PartnerService.ts', 'MEDIUM'],
            ['10', 'Video Knowledge Hub', 'New KnowledgeHub.tsx + admin content management', 'LOW'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Phase 4 - Polish & Scale (LOW Priority)', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Changes', 'Complexity'],
        [
            ['11', 'Live Activity Indicators', 'WebSocket enhancements (viewer count, online sellers)', 'LOW'],
            ['12', 'Phone OTP Authentication', 'Auth module enhancement + SMS provider (Twilio/MSG91)', 'MEDIUM'],
            ['13', 'Animal ID Card System', 'Animal module: breed-coded auto IDs + printable card', 'LOW'],
            ['14', 'Tipping / Platform Support', 'Wallet module enhancement + supporter badge', 'LOW'],
        ])
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════════════
    # 11. IMPACT SUMMARY
    # ═════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '11. Impact Summary', level=1)
    
    add_heading_styled(doc, 'Technical Impact', level=2)
    add_styled_table(doc,
        ['Metric', 'Before', 'After'],
        [
            ['Frontend Pages', '~50', '~60'],
            ['Backend Services', '44', '48+'],
            ['Database Tables', '~44', '~52'],
            ['API Routes', '55+', '70+'],
            ['Permission Keys', '~40', '~49'],
            ['Marketplace Categories', '7', '10+'],
            ['Nav Menu Items', '48', '55+'],
        ])
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Business Impact', level=2)
    biz_impacts = [
        'Revenue Streams: Boost listings (Silver/Gold tiers), marketplace transactions, partner commissions, subscription dairy deliveries',
        'User Engagement: Community forum creates sticky daily usage; milk recording creates daily login habit',
        'Trust & Verification: AI-verified milk yield, vaccination badges, IoT-verified weight, genomic profiles - all build buyer confidence',
        'Ecosystem Lock-in: Deep integration between modules means switching to a competitor loses all interconnected data',
        'Market Positioning: From "veterinary consultation app" to "complete livestock ecosystem platform" - competing directly with animall + pashushala while offering 10 exclusive innovations neither has',
        'Geographic Expansion: Multi-language support (5 languages) means immediate reach across India vs Hindi-only competitors',
    ]
    for impact in biz_impacts:
        add_bullet(doc, impact)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Competitive Advantage Summary', level=2)
    add_body_text(doc, 
        'VetCare will be the ONLY platform that combines veterinary consultation, livestock marketplace, '
        'dairy management, AI-powered recommendations, community forum, e-commerce (medicine + equipment + '
        'dairy products), and full traceability - all in a single integrated ecosystem with deep cross-module '
        'data flow. Neither animall.in nor pashushala.com comes close to this level of integration.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run('- End of Document -')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'
    run.italic = True
    
    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p2.add_run('VetCare Platform  |  Feature Analysis & Implementation Plan  |  March 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = 'Calibri'
    
    # ── Save ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 
                                'VetCare_Feature_Analysis_Plan.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    path = generate_document()
    print(f"\nSuccess! Word document generated at:\n{path}")
    print(f"File size: {os.path.getsize(path) / 1024:.1f} KB")
